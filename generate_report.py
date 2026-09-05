#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Trade Flow Analysis Report Generator
====================================

Generates a professional KEPROBA-style "Kenya – <Country> Trade Flow"
Word report (.docx) from:

  1. Seven Excel files downloaded from the ITC (International Trade Centre)
     database and manipulated into the standard format (see README):
       - "Figure 1 Trade Balance.xlsx"      bilateral trade over time
       - "Table 1 ... Import Source Markets.xlsx"
       - "Table 2 ... Lead Imports.xlsx"
       - "Table 3 ... Export destinations.xlsx"
       - "Table 4 ... Export Products.xlsx"
       - "Table 5 ... Exports to <Country>.xlsx"
       - "Table 6 ... imports from <Country>.xlsx"

  2. A per-country JSON configuration file that holds the country-specific
     content that is not in the Excel files (quick facts, world-trade
     context, optional images, references). The Section 1 "Backgrounds"
     narrative is auto-generated from the data; facts that require external
     research are printed as yellow-highlighted "[RESEARCH NEEDED: ...]"
     text for manual editing.

All trade tables, charts and the analytical narrative are computed
automatically from the Excel data.

Usage
-----
    python3 generate_report.py \
        --excel-dir sample_data \
        --config config/saudi_arabia.json \
        --output output/KENYA-SAUDI ARABIA TRADE FLOW.docx

Dependencies: python-docx, openpyxl, matplotlib
"""

import argparse
import json
import os
import re
import sys

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import openpyxl

import word_chart

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

from charts import draw_share_pie, series_shares, slice_callouts, new_fig, finish
from country_names import (display_name, is_africa, narrative_ref,
                           short_product_name, title_partner, title_case)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

YEAR_HEADER = [2021, 2022, 2023, 2024, 2025]

RED = RGBColor(0xC0, 0x00, 0x00)       # dark red used for Kenya rows
BRIGHT_RED = RGBColor(0xEE, 0x00, 0x00)  # red used for "Exports"/"Imports" labels
THEME_ACCENTS = ["#156082", "#E97132", "#196B24", "#0F9ED5", "#A02B93", "#4EA72E"]
# Century Gothic (the KEPROBA template font) is preferred; Liberation Serif is
# the metrically-similar Linux fallback when Century Gothic is not installed.
FONT_PREFERENCE = ["Century Gothic", "Liberation Serif", "DejaVu Serif", "Arial"]

SRC = ("Source: International Trade Centre (ITC) Trade Map database. "
       "Compiled by KEPROBA - Research and Innovation Directorate (RID).")


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------
def to_float(v):
    try:
        if v is None or v == "":
            return None
        f = float(v)
        return f if f == f else None
    except (TypeError, ValueError):
        return None


def num(v, decimals=1):
    """Format a number to `decimals` decimals with thousands separators,
    e.g. 3422.5 -> '3,422.5' (Word table/Excel-deliverable convention)."""
    if v is None:
        return ""
    n = f"{v:.{decimals}f}"
    if "." in n:
        whole, frac = n.split(".")
        frac = "." + frac
    else:
        whole, frac = n, ""
    sign = ""
    if whole.startswith("-"):
        sign, whole = "-", whole[1:]
    whole = f"{int(whole):,}" if whole.isdigit() else whole
    return sign + whole + frac


def pct(share, decimals=1):
    """Format a fraction as a percent, e.g. 0.27368 -> '27.4%'."""
    if share is None:
        return ""
    return f"{share * 100:.{decimals}f}%"


def _title_case(text):
    """Title-case a caption/title (delegates to the shared helper)."""
    return title_case(text)


def clean_label(label):
    """Remove the trailing '...' that ITC appends to long product labels."""
    if not label:
        return ""
    return re.sub(r"\s*\.\.\.\s*$", "", str(label)).rstrip(" ,;").strip()


def fit_widths(widths, n, prefix=3):
    """Rebuild a product-table column width list for `n` year columns.

    `widths` follows the template pattern
    ``[rank, code, label] + [value] * n_orig + [share]``; the value part is
    sliced to the actual number of years present in the data (repeating the
    last value if the template holds fewer year columns than the data).
    """
    if len(widths) == prefix + n + 1:
        return list(widths)
    share = widths[-1]
    vals = list(widths[prefix:-1])
    if vals and n > len(vals):
        vals = vals + [vals[-1]] * (n - len(vals))
    return list(widths[:prefix]) + vals[:n] + [share]


def short_label(label, maxlen=42):
    """Shorten a label for use inside narrative text.

    Delegates to the shared ``short_product_name`` so labels are cut at
    whole words (at least four kept) without any trailing '...'.
    """
    return short_product_name(label, maxlen=maxlen)


def lighten(hex_color, amount=0.55):
    """Mix a hex color with white."""
    hex_color = hex_color.lstrip("#")
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    mix = lambda c: int(round(c + (255 - c) * amount))
    return "#%02x%02x%02x" % (mix(r), mix(g), mix(b))


def resolve_path(base_file, path):
    """Resolve a path relative to the config file, falling back to the script dir."""
    if not path:
        return None
    p = os.path.expanduser(path)
    if os.path.isabs(p):
        return p
    candidate = os.path.join(os.path.dirname(os.path.abspath(base_file)), p)
    if os.path.exists(candidate):
        return candidate
    candidate = os.path.join(BASE_DIR, p)
    return candidate if os.path.exists(candidate) else None


def chart_font():
    """Pick the first installed font from the preference list (cross-platform)."""
    import matplotlib.font_manager as fm
    installed = {f.name for f in fm.fontManager.ttflist}
    for name in FONT_PREFERENCE:
        if name in installed:
            return name
    return fm.fontManager.defaultFont["family"]


# ---------------------------------------------------------------------------
# Excel reading
# ---------------------------------------------------------------------------
def find_excel_files(excel_dir):
    """Locate the seven Excel files by keyword in their filenames."""
    found = {}
    for fname in sorted(os.listdir(excel_dir)):
        low = fname.lower()
        if not low.endswith((".xlsx", ".xlsm")):
            continue
        path = os.path.join(excel_dir, fname)
        if "table 1" in low:
            found["table1"] = path
        elif "table 2" in low:
            found["table2"] = path
        elif "table 3" in low:
            found["table3"] = path
        elif "table 4" in low:
            found["table4"] = path
        elif "table 5" in low:
            found["table5"] = path
        elif "table 6" in low:
            found["table6"] = path
        elif "table 7" in low:
            found["table7"] = path          # optional: bilateral alignment
        elif "product by partner" in low or "table 8" in low:
            found["product_partner"] = path  # optional: per-product origin detail
        elif "balance" in low or "figure 1" in low:
            found["balance"] = path
    missing = [k for k in
               ("table1", "table2", "table3", "table4", "table5", "table6", "balance")
               if k not in found]
    if missing:
        expected = {
            "table1": "Table 1",
            "table2": "Table 2",
            "table3": "Table 3",
            "table4": "Table 4",
            "table5": "Table 5",
            "table6": "Table 6",
            "balance": "Balance or Figure 1",
        }
        missing_detail = []
        for k in missing:
            missing_detail.append(f"  - {k}: expected filename containing '{expected[k]}'")
        found_names = {k: os.path.basename(v) for k, v in found.items()}
        sys.exit(
            "[ERROR] Missing Excel file(s) in '%s'.\n"
            "Found %d file(s): %s\n"
            "Missing %d required file(s):\n%s"
            % (excel_dir, len(found),
               ", ".join(f"{k}={v}" for k, v in sorted(found_names.items())),
               len(missing), "\n".join(missing_detail)))
    return found


def _cell_grid(ws):
    return [[c.value for c in row] for row in ws.iter_rows()]


def _find_year_run(grid):
    """Return (row_index, start_col, [years]) for the longest run of
    consecutive year integers (between 2 and 10 years) anywhere in the grid,
    else (None, None, []).

    This lets the parser accept any consecutive year range (e.g. 2021..2025,
    2021..2024, 2015..2024 or a short 2023..2025 window as produced by the
    classic Trade Map interface), not just a fixed five-year block.
    """
    best = None
    for ri, row in enumerate(grid):
        cur = []
        prev = None
        for ci, v in enumerate(row):
            iv = None
            if isinstance(v, bool):
                iv = None
            elif isinstance(v, (int, float)):
                iv = int(v)
            elif isinstance(v, str):
                s = v.strip().strip("'\"").replace(",", "")
                if re.fullmatch(r"\d{4}", s) and 1990 <= int(s) <= 2100:
                    iv = int(s)
            if iv is not None and 1990 <= iv <= 2100 and (not isinstance(v, str) and abs(v - iv) < 1e-9
                                                          or isinstance(v, str)):
                if prev is not None and iv == prev + 1:
                    cur.append((ci, iv))
                else:
                    if 2 <= len(cur) <= 10 and (best is None or len(cur) > len(best[2])):
                        best = (ri, cur[0][0], [y for _, y in cur])
                    cur = [(ci, iv)]
                prev = iv
            else:
                prev = None
        if 2 <= len(cur) <= 10 and (best is None or len(cur) > len(best[2])):
            best = (ri, cur[0][0], [y for _, y in cur])
    return best if best else (None, None, [])


def parse_rank_table(path):
    """
    Parse a generic ITC "rank table" (Tables 1-6).

    The layout detected is:
        ... header rows ...
        ... [optional header cell containing the word "code"] ...
        ... a row with the consecutive years (any run of 4 to 10 consecutive
            years, e.g. 2021..2025 or 2021..2024) ...
        [optional row: "Value in USD Billion"]
        rank | [code] | name/label | y1 .. yn | share(fraction)
        ...
        "All other ..."  /  "World" | "All products"
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    grid = _cell_grid(ws)

    year_row_idx, year_start_col, year_header = _find_year_run(grid)
    if year_row_idx is None:
        raise ValueError(f"Could not find a consecutive year header row in {path}")
    n = len(year_header)

    val_cols = list(range(year_start_col, year_start_col + n))
    share_col = year_start_col + n

    has_code = False
    for ri in range(year_row_idx + 1):
        for v in grid[ri]:
            if isinstance(v, str) and "code" in v.strip().strip("'\"`").lower():
                has_code = True
                break
        if has_code:
            break

    name_col = 2 if has_code else 1

    data = []
    for ri in range(year_row_idx + 1, len(grid)):
        row = grid[ri]
        if not any(v is not None and str(v).strip() != "" for v in row):
            continue
        rank = int(row[0]) if isinstance(row[0], (int, float)) else None
        name = row[name_col] if name_col < len(row) else None
        code = row[1] if has_code and len(row) > 1 else None
        label = row[2] if has_code and len(row) > 2 else None
        if name is None and label is not None:
            name = label
        if name is None or (isinstance(name, str) and not name.strip()):
            continue
        name = str(name).strip()
        years = [to_float(row[c]) if c < len(row) else None for c in val_cols]
        share = to_float(row[share_col]) if share_col < len(row) else None

        if name.lower().startswith("all other"):
            kind = "all_other"
        elif name.lower() in ("world", "all products"):
            kind = "total"
        else:
            kind = "item"
        data.append({"rank": rank, "name": name, "code": code, "label": label,
                     "years": years, "share": share, "kind": kind})

    items = [d for d in data if d["kind"] == "item"]
    total = next((d for d in data if d["kind"] == "total"), None)
    all_other = next((d for d in data if d["kind"] == "all_other"), None)
    return {"has_code": has_code, "data": data, "items": items,
            "total": total, "all_other": all_other,
            "year_start_col": year_start_col, "years": year_header}


def parse_product_partner(path):
    """Parse an optional ITC "trade by product and partner" download into
    rows of {product_code, partner, value_usd}.

    ITC structurally varies this export across interfaces; we keep the parser
    tolerant: scan the header row for a product-code column, a partner-country
    column and a numeric trade-value column, then emit one row per line.
    Returns [] when the shape is not recognised (caller treats it as optional).
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    grid = _cell_grid(ws)
    rows = []
    prod_col = partner_col = value_col = None
    # find a header row (a row with >=2 string cells mentioning code/partner/value)
    for ri, row in enumerate(grid[:40]):
        cells = [str(c).strip().lower() for c in row if isinstance(c, str)]
        joined = " ".join(cells)
        has_code = any("code" in c or "hs" in c or "product" in c for c in cells)
        has_part = any("partner" in c or "country" in c or "report" in c for c in cells)
        has_val = any(("value" in c or "usd" in c or "amount" in c) for c in cells)
        if has_code and has_val and has_part:
            for ci, c in enumerate(row):
                if not isinstance(c, str):
                    continue
                cc = c.strip().lower()
                if prod_col is None and ("product code" in cc or "hs code" in cc
                                         or cc.endswith("hs") or cc == "commodity code"):
                    prod_col = ci
                elif partner_col is None and ("partner" in cc
                                              or ("country" in cc and "report" not in cc)):
                    partner_col = ci
                elif value_col is None and ("trade value" in cc or "import value"
                                            in cc or cc.endswith("usd") or "value (usd" in cc
                                            or "value usd" in cc or cc == "value"):
                    value_col = ci
            for rj in range(ri + 1, len(grid)):
                r = grid[rj]
                if not r:
                    continue
                if prod_col is None or prod_col >= len(r):
                    continue
                code = r[prod_col]
                partner = r[partner_col] if (partner_col is not None and partner_col < len(r)) else None
                val = None
                try:
                    # value may be spread over several year columns; take max = latest
                    vals = []
                    for c in range(value_col if value_col is not None else 0, len(r)):
                        if isinstance(r[c], (int, float)):
                            vals.append(float(r[c]))
                    if vals:
                        val = max(vals)
                except Exception:
                    val = None
                if code is None or val is None:
                    continue
                rows.append({"product_code": str(code).replace("'", "").strip(),
                             "partner": str(partner).strip() if partner else "",
                             "value_usd": val})
            break
    return rows


def parse_trade_balance(path):
    """Parse the 'Figure 1 Trade Balance' workbook (bilateral exports/imports)."""
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    grid = _cell_grid(ws)

    for ri, row in enumerate(grid):
        if not row or not isinstance(row[0], str):
            continue
        lbl = row[0].strip()
        if lbl == "Exports" and ri + 2 < len(grid):
            l1 = grid[ri + 1][0].strip() if isinstance(grid[ri + 1][0], str) else ""
            l2 = grid[ri + 2][0].strip() if isinstance(grid[ri + 2][0], str) else ""
            if l1 == "Imports" and l2 == "Balance of Trade":
                years = [int(v) for v in grid[ri - 1][1:] if isinstance(v, (int, float)) and int(v) >= 2000]
                if not years:
                    continue
                exports = [to_float(row[1 + i]) for i in range(len(years))]
                imports = [to_float(grid[ri + 1][1 + i]) for i in range(len(years))]
                balance = [to_float(grid[ri + 2][1 + i]) for i in range(len(years))]
                return {"years": years, "exports": exports, "imports": imports, "balance": balance}
    raise ValueError(f"Could not parse bilateral trade series from {path}")


def parse_bilateral_table(path):
    """Parse the optional Table 7 workbook (export composition & alignment).

    Layout as written by ``make_tables.write_bilateral_table``:
        row 1: "Table 7:" | merged title
        row 2: Rank | HS Code | Product | Kenya Value USD M (YYYY) |
               Kenya Share (%) | <Partner> Import Share (%) | Annual Growth %
        rows 3+: one row per product; shares/growth stored as fractions.
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    try:
        ws = wb.worksheets[0]
        grid = _cell_grid(ws)
    finally:
        wb.close()

    hdr_idx = None
    year = None
    for ri, row in enumerate(grid[:5]):
        for v in row[:4]:
            if isinstance(v, str):
                m = re.search(r"\((20\d{2})\)", v)
                if m:
                    year, hdr_idx = int(m.group(1)), ri
                    break
        if hdr_idx is not None:
            break
    if hdr_idx is None:
        raise ValueError(f"Could not locate the Table 7 header row in {path}")

    items = []
    for row in grid[hdr_idx + 1:]:
        if len(row) < 7:
            continue
        rank = row[0]
        if not isinstance(rank, (int, float)) or isinstance(rank, bool):
            continue
        label = str(row[2]).strip() if row[2] is not None else ""
        if not label:
            continue
        items.append({
            "rank": int(rank),
            "code": str(row[1]).strip() if row[1] is not None else "",
            "label": clean_label(label),
            "value": to_float(row[3]),
            "kenya_share": to_float(row[4]),
            "partner_share": to_float(row[5]),
            "growth": to_float(row[6]),
        })
    return {"year": year, "items": items}


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------
class Analysis:
    """Holds every statistic needed by the narrative and the charts."""

    def __init__(self, cfg):
        self.cfg = cfg
        self.country = cfg["country"]["name"]
        self.short = cfg["country"]["short"]
        self.possessive = cfg["country"]["possessive"]
        self.year = int(cfg["report"]["year"])
        self.years = YEAR_HEADER
        self.iy = 0                       # set in load(), once the data is read
        self.i5 = 0

    def load(self, excel_dir):
        files = find_excel_files(excel_dir)
        self.table1 = parse_rank_table(files["table1"])   # import source markets
        self.table2 = parse_rank_table(files["table2"])   # import products
        self.table3 = parse_rank_table(files["table3"])   # export destinations
        self.table4 = parse_rank_table(files["table4"])   # export products
        self.table5 = parse_rank_table(files["table5"])   # Kenya exports to country
        self.table6 = parse_rank_table(files["table6"])   # Kenya imports from country
        self.balance = parse_trade_balance(files["balance"])
        try:
            self.table7 = (parse_bilateral_table(files["table7"])
                           if "table7" in files else None)
        except Exception:
            self.table7 = None
        try:
            if "product_partner" in files:
                pp = parse_product_partner(files["product_partner"])
                self.set_africa_import_share(self.load_africa_import_share(pp))
        except Exception:
            self.set_africa_import_share({})
        if self.table1["years"]:
            self.years = self.table1["years"]
        if self.year not in self.years:
            self.year = self.years[-1]     # e.g. 2024 when the data stops in 2024
        self.iy = self.years.index(self.year)
        self.i5 = len(self.years) - 1
        self._warn_hs_level()
        return self

    def _warn_hs_level(self):
        """Heads-up when a product table appears to be at HS2 rather than the
        required HS4 detail (ITC's Trade Map defaults to a coarser or finer
        level depending on how the download was configured).  This is a
        warning only -- the report still builds and the tables remain valid."""
        def _prod_codes(parsed):
            return [str(d.get("code") or "") for d in (parsed or {}).get("items", [])]

        def _digits(codes):
            digs = [len(c.strip().strip("'"))
                    for c in codes if c.strip().strip("'").isdigit()]
            return digs[0] if digs else None

        for tag, parsed in (("Table 4 (export products)", self.table4),
                            ("Table 5 (exports to market)", self.table5),
                            ("Table 6 (imports from market)", self.table6)):
            lvl = _digits(_prod_codes(parsed))
            if lvl is not None and lvl != 4:
                why = "" if lvl in (2, 6) else \
                    f" (looks like a mix/malformed: HS{lvl})"
                print(f"[WARN] {tag}: product codes look like HS{lvl}"
                      f"{why}. This report requires HS4 (4-digit Harmonized "
                      f"System) product detail for consistent analysis. "
                      f"Re-download from ITC Trade Map at 4-digit level for "
                      f"the correct product granularity.",
                      file=sys.stderr)

    # ---- country aggregates ------------------------------------------------
    def imports_years(self):
        return self.table1["total"]["years"] if self.table1["total"] else [None] * len(self.years)

    def exports_years(self):
        return self.table3["total"]["years"] if self.table3["total"] else [None] * len(self.years)

    def kenya_exports_years(self):
        return self.table5["total"]["years"] if self.table5["total"] else [None] * len(self.years)

    def kenya_imports_years(self):
        return self.table6["total"]["years"] if self.table6["total"] else [None] * len(self.years)

    def imports_2025(self):
        return self.imports_years()[self.iy]

    def imports_2024(self):
        return self.imports_years()[self.iy - 1]

    def imports_2021(self):
        return self.imports_years()[0]

    def exports_2025(self):
        return self.exports_years()[self.iy]

    def exports_2024(self):
        return self.exports_years()[self.iy - 1]

    def imports_growth_2024_25(self):
        a, b = self.imports_2024(), self.imports_2025()
        return (b / a - 1) if a and b else None

    def imports_cagr_2021_25(self):
        a, b = self.imports_2021(), self.imports_2025()
        n = len(self.years)
        return ((b / a) ** (1.0 / (n - 1)) - 1) if a and b and n > 1 else None

    def exports_growth_2024_25(self):
        a, b = self.exports_2024(), self.exports_2025()
        return (b / a - 1) if a and b else None

    def exports_cagr(self):
        """Compound annual growth of Kenya's bilateral exports across the
        full series present in the data (may be more than 5 years)."""
        yrs = self.kenya_exports_years()
        vals = [v for v in yrs if v is not None]
        if len(vals) < 2 or vals[0] in (0, None) or vals[-1] is None:
            return None
        n = len(yrs) - 1
        return (vals[-1] / vals[0]) ** (1.0 / n) - 1 if n else None

    def top(self, table, n):
        return table["items"][:n]

    def top10_export_share(self):
        vals = [d["share"] for d in self.table5["items"][:10]]
        vals = [v for v in vals if v is not None]
        return sum(vals) if vals else None

    def top10_import_share(self):
        vals = [d["share"] for d in self.table6["items"][:10]]
        vals = [v for v in vals if v is not None]
        return sum(vals) if vals else None

    @staticmethod
    def _hhi(shares):
        """Herfindahl-Hirschman Index over a set of (normalised-to-1) shares.
        Returns None when there is insufficient data."""
        shares = [s for s in shares if s is not None and s > 0]
        if not shares:
            return None
        return sum(s * s for s in shares)

    def exports_hhi(self, items=None):
        """HHI of the partner market's top export-product share basket
        (Table 4) in the latest year."""
        items = items if items is not None else self.table4["items"]
        return self._hhi([d.get("share") for d in items])

    def imports_hhi(self, items=None):
        """HHI of the partner market's top import-product share basket
        (Table 2) in the latest year."""
        items = items if items is not None else self.table2["items"]
        return self._hhi([d.get("share") for d in items])

    def kenya_exports_concentration(self):
        """HHI of Kenya's bilateral export-product shares to the partner
        (Table 5) - a measure of how concentrated Kenya's exports are."""
        items = self.table5["items"]
        shares = [d.get("share") for d in items]
        if all(s is None for s in shares):
            return None
        return self._hhi(shares)

    def kenya_imports_concentration(self):
        """HHI of Kenya's bilateral import-product shares from the partner
        (Table 6)."""
        items = self.table6["items"]
        shares = [d.get("share") for d in items]
        if all(s is None for s in shares):
            return None
        return self._hhi(shares)

    @staticmethod
    def _cagr(items):
        """Compound annual growth over a full series of yearly values."""
        vals = [v for v in items if v is not None]
        if len(vals) < 2 or not vals[0] or not vals[-1]:
            return None
        n = len(items) - 1
        return (vals[-1] / vals[0]) ** (1.0 / n) - 1 if n else None

    def exports_product_cagr(self):
        return self._cagr(self.kenya_exports_years())

    def imports_product_cagr(self):
        return self._cagr(self.kenya_imports_years())

    def kenya_exports_5yr(self):
        return self.kenya_exports_years()

    def kenya_exports_average(self):
        vals = [v for v in self.kenya_exports_5yr() if v is not None]
        return sum(vals) / len(vals) if vals else None

    def kenya_exports_max(self):
        vals = self.kenya_exports_5yr()
        if not vals:
            return None, None
        idx = max(range(len(vals)), key=lambda i: vals[i] if vals[i] is not None else -1e18)
        return vals[idx], self.years[0] + idx

    def balance_always_negative(self):
        b = self.balance["balance"]
        return bool(b) and all(v is not None and v < 0 for v in b)

    def balance_imports_fluct(self):
        """Highs/lows of the full bilateral import series from the balance file.

        Returns {max, max_year, min, min_year, next_growth, next_year} where
        next_growth is the % change from the minimum year to the following year
        (None when the minimum is the last available year).
        """
        ys = self.balance["years"]
        imps = self.balance["imports"]
        if not imps:
            return None
        mx_i = max(range(len(imps)), key=lambda i: imps[i] if imps[i] is not None else -1e18)
        mn_i = min(range(len(imps)), key=lambda i: imps[i] if imps[i] is not None else 1e18)
        res = {"max": imps[mx_i], "max_year": ys[mx_i],
               "min": imps[mn_i], "min_year": ys[mn_i]}
        if mn_i + 1 < len(imps) and imps[mn_i + 1] and imps[mn_i]:
            res["next_growth"] = imps[mn_i + 1] / imps[mn_i] - 1
            res["next_year"] = ys[mn_i + 1]
        else:
            res["next_growth"] = None
        return res

    def country_top_markets_title_n(self, table):
        """Count of ranked partner rows (excluding the Kenya row)."""
        n = 0
        for d in table["items"]:
            if d["rank"] is not None and d["name"].lower() != "kenya":
                n += 1
        return n

    # ---- Africa / Kenya destination focus (from the partner's export market
    #      table, Table 3) ----------------------------------------------------
    def africa_destinations(self, table=None):
        """Rows of `table` (Table 3 export destinations by default) whose
        destination is an African country.  Returns the subset of the items
        list (so callers can reuse .name / .share / .years)."""
        table = table or self.table3
        return [d for d in table["items"] if is_africa(d["name"])]

    def africa_export_share(self, table=None):
        """Share (fraction 0..1) of the partner's exports destined to Africa in
        the latest year, summed over the African destination rows present."""
        rows = self.africa_destinations(table)
        vals = [d["share"] for d in rows if d.get("share") is not None]
        return sum(vals) if vals else None

    def kenya_destination(self, table=None):
        """The Kenya row in the partner's export-destination table, if any."""
        table = table or self.table3
        for d in table["items"]:
            if is_africa(d["name"]) and d["name"].lower() == "kenya":
                return d
        return None

    def kenya_destination_share(self, table=None):
        k = self.kenya_destination(table)
        return k.get("share") if k else None

    def africa_import_sources(self, table=None):
        """Rows of `table` (Table 1 import source markets) whose origin is an
        African country."""
        table = table or self.table1
        return [d for d in table["items"] if is_africa(d["name"])]

    def kenya_import_source(self, table=None):
        """The Kenya row in the partner's import source-market table (Table 1)."""
        table = table or self.table1
        for d in table["items"]:
            if is_africa(d["name"]) and d["name"].lower() == "kenya":
                return d
        return None

    def africa_import_share(self, table=None):
        """Share (fraction 0..1) of the partner's imports sourced from Africa
        groups (Table 1) in the latest year."""
        table = table or self.table1
        vals = [d["share"] for d in self.africa_import_sources(table)
                if d.get("share") is not None]
        return sum(vals) if vals else None

    # ---- Top-N imported products: worldwide vs Africa vs Kenya -------------
    @staticmethod
    def _hs(item):
        """Normalise an item's HS code to a bare str (dropping a leading ')."""
        c = str((item or {}).get("code") or "").replace("'", "").strip()
        return c

    def top_import_products(self, n=20):
        """For the partner market's top ``n`` imported products (Table 2,
        worldwide, USD billion) return a list of dicts per product with:

          code, name,        HS4 code + label
          worldwide_val_bin  partner's total world imports of the product (USD bn)
          kenya_val_bin      partner's imports of the product from Kenya (USD bn,
                             from Table 5 which is in USD million -- converted)
          kenya_share        Kenya's share (0..1) of that product's world imports

        ``africa_share`` is only populated when a per-product source-partner
        breakdown is supplied (see ``africa_import_products``); otherwise None.
        """
        t2 = {self._hs(d): d for d in self.table2["items"]}
        t5 = {self._hs(d): d for d in self.table5["items"]}
        out = []
        for d in self.table2["items"][:n]:
            c = self._hs(d)
            ww = d["years"][-1] if d.get("years") else None          # USD bn
            kd = t5.get(c)
            kv = (kd["years"][-1] if kd and kd.get("years") else None)  # USD mn
            kv_bn = (kv / 1000.0) if kv is not None else None        # -> USD bn
            kshare = None
            if kv_bn is not None and ww:
                kshare = kv_bn / ww
            out.append({
                "rank": d.get("rank"),
                "code": c,
                "name": d.get("name") or d.get("label"),
                "worldwide_val_bin": ww,
                "kenya_val_bin": kv_bn,
                "kenya_share": kshare,
                "africa_share": None,
            })
        if hasattr(self, "_africa_import_share"):
            for row in out:
                row["africa_share"] = self._africa_import_share.get(row["code"])
        return out

    def set_africa_import_share(self, mapping):
        """Attach per-HS4 ``{code: africa_share}`` (0..1) so ``top_import_products``
        can report the 'imports from Africa' share.  Provided by an optional
        product-by-partner breakdown (Trade Map download)."""
        self._africa_import_share = mapping or {}
        return self

    # ---- Optional per-product Africa import origin ------------------------
    def load_africa_import_share(self, product_partner_items):
        """From an optional ITC 'trade by product and partner' breakdown
        (list of {product_code, partner, value_usd}), compute the Africa share
        of the partner's world imports per HS4 product."""
        africa = {}
        world = {}
        for row in product_partner_items or []:
            code = str(row.get("product_code") or "").replace("'", "").strip()[:4]
            val = row.get("value_usd")
            if not code or val is None:
                continue
            world[code] = world.get(code, 0.0) + val
            if is_africa(str(row.get("partner") or "")):
                africa[code] = africa.get(code, 0.0) + val
        mapping = {}
        for code, w in world.items():
            a = africa.get(code, 0.0)
            mapping[code] = (a / w) if w else 0.0
        return mapping


# ---------------------------------------------------------------------------
# Narrative text generation
# ---------------------------------------------------------------------------
def build_narratives(a: Analysis, cfg):
    c = cfg["country"]
    Y = a.year
    t = {}
    world = cfg.get("world_trade", {}) or {}

    _NEXT_UNIT = {"thousand": "million", "million": "billion", "billion": "trillion"}

    def usd_auto(v, dec=1, unit="billion"):
        """Format v (expressed in `unit`) upgrading to the next unit at 1,000,
        e.g. 2,200 billion -> "2.2 trillion", 1,500 million -> "1.5 billion".
        The Excel files keep the raw unit; only the narrative upgrades."""
        if v is None:
            return ""
        if v >= 1000 and unit in _NEXT_UNIT:
            return f"{v / 1000:,.{dec}f} {_NEXT_UNIT[unit]}"
        return f"{v:,.{dec}f} {unit}"

    n_years_words = {2: "two", 3: "three", 4: "four", 5: "five",
                     6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten"}
    wn = n_years_words.get(len(a.years), str(len(a.years)))

    # ---- world-trade reference values (used by sections 1, 2) -------------
    wex = world.get("world_exports_usd_billion")
    wim = world.get("world_imports_usd_billion")
    exp_2025, imp_2025 = a.exports_2025(), a.imports_2025()

    # ---- Section 1 background ---------------------------------------------
    # Follows the reference report's four-paragraph flow (overview, economy,
    # outlook, trade & policy). Sentences that can be computed from the Excel
    # data are generated here; facts that must be researched externally are
    # returned as flagged segments (rendered highlighted) for manual editing.
    def mkt_summary(rows):
        return top_phrase(rows, lambda d: f"{d['name']} ({pct(d['share'])})")

    def top_prods(rows):
        return top_phrase(rows, lambda d: short_label(d["label"]) or d["name"])

    def research(what):
        return f"[RESEARCH NEEDED: {what}]"

    def top_phrase(items, fmt, n=None):
        """Join up to n top items as "A, B, and C", tolerating fewer rows
        (a small country table may have only a handful of product lines)."""
        rows = items if n is None else items[:n]
        if not rows:
            return ""
        names = [fmt(d) for d in rows]
        if len(names) == 1:
            return names[0]
        if len(names) == 2:
            return f"{names[0]} and {names[1]}"
        return ", ".join(names[:-1]) + f", and {names[-1]}"

    wex_s = f", equivalent to {exp_2025 / wex * 100:.1f}% of world exports" if wex else ""
    wim_s = f", or {imp_2025 / wim * 100:.1f}% of world imports" if wim else ""

    t["background"] = [
        [  # 1) country overview
            (f"{c['name']} has its capital at {c['capital']}. ", False),
            (research("country overview - geography, total area, borders, population "
                      "size and growth, demographics, urbanisation, life expectancy"), True),
        ],
        [  # 2) economy
            (research("economy - GDP and GDP per capita (current USD), real GDP growth, "
                      "key economic sectors, income classification (e.g. high-income)"), True),
        ],
        [  # 3) outlook
            (research("outlook - IMF World Economic Outlook and World Bank growth/inflation "
                      "projections, fiscal and public-debt position, policy priorities"), True),
        ],
        [  # 4) trade and policy
            (f"On trade, {c['name']} exported goods worth USD {usd_auto(exp_2025)} in {Y}{wex_s}, "
             f"while imports were valued at USD {usd_auto(imp_2025)}{wim_s}. ", False),
            (f"The leading import source markets were {mkt_summary(a.top(a.table1, 3))}, "
             f"and the principal export destinations were {mkt_summary(a.top(a.table3, 3))}. ", False),
            (f"The top import products were {top_prods(a.top(a.table2, 3))}, while the top export "
             f"products were {top_prods(a.top(a.table4, 3))}. ", False),
            (research("trade and development policy - WTO accession date, main trade agreements, "
                      "national development/export-diversification strategy"), True),
        ],
    ]

    # ---- Section 2 intro --------------------------------------------------
    exp_world = f", equivalent to {exp_2025 / wex * 100:.1f}% of world exports" if wex else ""
    imp_world = f", or {imp_2025 / wim * 100:.1f}% of world imports" if wim else ""

    dest1 = a.top(a.table3, 1)[0]
    prod1 = a.top(a.table4, 1)[0]
    src1 = a.top(a.table1, 1)[0]
    imp1 = a.top(a.table2, 1)[0]

    t["exports_bullets"] = [
        f"{c['name']} exports in {Y} were valued at USD {usd_auto(exp_2025)}{exp_world}.",
        f"The lead export destination market in {Y} was {dest1['name']}, with exports of "
        f"USD {usd_auto(dest1['years'][a.iy])}. This represented {pct(dest1['share'])} of "
        f"{a.possessive} total exports.",
        f"The principal export product was {clean_label(prod1['label']) or prod1['name']}, "
        f"valued at USD {usd_auto(prod1['years'][a.iy])}. It ranked first among "
        f"{c['name']}'s export products.",
    ]
    t["imports_bullets"] = [
        f"In {Y}, {c['name']} imports were valued at USD {usd_auto(imp_2025)}{imp_world}.",
        f"The main import product was {clean_label(imp1['label']) or imp1['name']}, worth "
        f"USD {usd_auto(imp1['years'][a.iy])}. It accounted for {pct(imp1['share'])} of "
        f"the country's imports.",
        f"The leading import source market was {src1['name']}, which supplied products worth "
        f"USD {usd_auto(src1['years'][a.iy])}. This represented {pct(src1['share'])} of "
        f"{c['name']} imports.",
    ]

    # ---- Section 2.1 ------------------------------------------------------
    growth = a.imports_growth_2024_25()
    cagr = a.imports_cagr_2021_25()
    s21_mkts = top_phrase(
        a.top(a.table1, 3),
        lambda d: f"{d['name']} (USD {usd_auto(d['years'][a.iy])}; {pct(d['share'])})")
    src3 = a.top(a.table1, 3)
    top3_import_share = sum((d.get("share") or 0) for d in src3)
    s21_analysis = (
        f"The three leading source markets together supplied {pct(top3_import_share)} "
        f"of total imports. This concentration indicates that supply is heavily "
        f"dependent on a few partners, so diversifying import sources would reduce risk."
        if top3_import_share and top3_import_share >= 25 else "")
    t["s21"] = [
        f"In {Y}, {a.possessive} imports were valued at USD {usd_auto(a.imports_2025())}, "
        f"an increase of {pct(growth)} on the previous year's USD {usd_auto(a.imports_2024())}.",
        f"Between {a.years[0]} and {a.years[-1]}, imports grew at an average of "
        f"{round(cagr * 100)}% a year." if cagr else "",
        f"The lead source markets in {Y} were {s21_mkts}." if s21_mkts else "",
        s21_analysis if s21_analysis else "",
    ]

    # ---- Section 2.1a: Africa / peer benchmark of import sources ----------
    afr_imp_share = a.africa_import_share()
    kenya_src = a.kenya_import_source()
    src_benchmark = []
    if afr_imp_share is not None:
        src_benchmark.append(
            f"Of {a.country}'s total goods imports in {Y}, {pct(afr_imp_share)} "
            f"originated from Africa.")
    if kenya_src and kenya_src.get("share") is not None:
        rank_txt = (f" (ranking {kenya_src['rank']} among source markets)"
                    if kenya_src.get("rank") else "")
        src_benchmark.append(
            f"Kenya supplied {pct(kenya_src['share'])} of the country's imports{rank_txt}.")
    t["s21_kenya"] = src_benchmark


    # ---- Section 2.2 ------------------------------------------------------
    s22_prods = top_phrase(
        a.top(a.table2, 4),
        lambda d: f"{clean_label(d['label'])} (USD {usd_auto(d['years'][a.iy])}, {pct(d['share'])})")
    t["s22"] = [
        f"In {Y}, the leading import products were: {s22_prods}." if s22_prods else "",
    ]

    # ---- Section 2.2a: Kenya's share of the partner's top imports --------
    ke_insights = []
    kimp = a.top_import_products(n=6)
    kimp_rows = [r for r in kimp if r.get("kenya_share") is not None]
    if kimp_rows:
        best = max(kimp_rows, key=lambda r: r["kenya_share"])
        ks = top_phrase(
            sorted(kimp_rows, key=lambda r: r["kenya_share"], reverse=True)[:3],
            lambda r: f"{clean_label(r['name'])} ({pct(r['kenya_share'])})")
        ke_insights.append(
            f"Within the partner's top imported products, Kenya's largest shares were "
            f"in {ks}.")
        gap = best["kenya_share"]
        ke_insights.append(
            f"Kenya's single biggest foothold was in {clean_label(best['name'])} at "
            f"{pct(gap)} of that product's world imports, showing room to scale "
            f"bilateral supply.")
    afr_rows = [r for r in kimp if r.get("africa_share") is not None]
    if afr_rows:
        top_afr = max(afr_rows, key=lambda r: r["africa_share"])
        if top_afr.get("africa_share") is not None and top_afr.get("kenya_share") is not None:
            ke_insights.append(
                f"Africa supplies {pct(top_afr['africa_share'])} of world imports of "
                f"{clean_label(top_afr['name'])}, of which Kenya provides "
                f"{pct(top_afr['kenya_share'])}, pointing to a broader regional "
                f"sourcing opportunity.")
    t["s22_kenya"] = ke_insights

    # ---- Section 2.3 ------------------------------------------------------
    eg = a.exports_growth_2024_25()
    s23_dsts = top_phrase(
        a.top(a.table3, 4),
        lambda d: f"{d['name']} (USD {usd_auto(d['years'][a.iy])}, {pct(d['share'])})")
    dest3 = a.top(a.table3, 3)
    top3_export_share = sum((d.get("share") or 0) for d in dest3)
    s23_analysis = (
        f"The three leading destinations took {pct(top3_export_share)} of total exports. "
        f"Such concentration means export earnings depend heavily on a small set of "
        f"markets, underlining the value of market diversification."
        if top3_export_share and top3_export_share >= 25 else "")
    afr_share = a.africa_export_share()
    kenya_dest = a.kenya_destination()
    afr_bullet = ""
    if afr_share is not None:
        afr_countries = ", ".join(
            display_name(d["name"]) for d in a.africa_destinations() if d.get("share"))
        msg = (f"Of {a.country}'s total exports in {Y}, "
               f"{pct(afr_share)} were destined to Africa"
               + (f" ({afr_countries})." if afr_countries else "."))
        if kenya_dest and kenya_dest.get("share") is not None:
            rank_txt = (f" (ranking {kenya_dest['rank']} among destinations)"
                        if kenya_dest.get("rank") else "")
            msg += (f" Exports to Kenya accounted for {pct(kenya_dest['share'])}"
                    f" of the total{rank_txt}.")
        afr_bullet = msg
    t["s23"] = [
        (f"In {Y}, {a.possessive} exports were valued at USD {usd_auto(a.exports_2025())}, "
         f"up {round(eg * 100)}% from USD {usd_auto(a.exports_2024())} in {Y - 1}."
         if eg else ""),
        f"The leading destination markets were {s23_dsts}." if s23_dsts else "",
        afr_bullet if afr_bullet else "",
        s23_analysis if s23_analysis else "",
    ]

    # ---- Section 2.4 ------------------------------------------------------
    s24_prods = top_phrase(
        a.top(a.table4, 5),
        lambda d: f"{clean_label(d['label'])} (USD {usd_auto(d['years'][a.iy])}, {pct(d['share'])})")
    t["s24"] = [
        f"In {Y}, {a.possessive} top export products were: {s24_prods}." if s24_prods else "",
    ]

    # ---- Section 3.1 ------------------------------------------------------
    ke = a.kenya_exports_2025 = a.kenya_exports_years()[a.iy]
    avg = a.kenya_exports_average()
    mx, mxy = a.kenya_exports_max()
    t["s31"] = [
        f"In {Y}, Kenya's exports to {c['name']} were valued at USD {usd_auto(ke, unit='million')}.",
        f"Across {a.years[0]}-{a.years[-1]}, exports to {c['name']} averaged "
        f"USD {usd_auto(avg, unit='million')} a year. The peak value was "
        f"USD {usd_auto(mx, unit='million')}, recorded in {mxy}.",
        (f"Exports in {Y} stood {('above' if ke >= avg else 'below')} the {wn}-year average, "
         f"indicating a {('favourable' if ke >= avg else 'softening')} short-term trend."
         if avg else ""),
        f"In {Y}, Kenya's imports from {c['name']} were valued at "
        f"USD {usd_auto(a.kenya_imports_years()[a.iy], unit='million')}.",
    ]
    fluct = a.balance_imports_fluct()
    if fluct:
        msg = (f"Imports from {c['name']} have fluctuated markedly. "
               f"The highest value was USD {usd_auto(fluct['max'], unit='million')}, "
               f"recorded in {fluct['max_year']}.")
        if fluct.get("next_growth") is not None:
            msg += (f" After a low of USD {usd_auto(fluct['min'], unit='million')} in "
                    f"{fluct['min_year']}, imports rose by {fluct['next_growth'] * 100:.1f}% "
                    f"in {fluct['next_year']}.")
        else:
            msg += (f" The lowest value was USD {usd_auto(fluct['min'], unit='million')}, "
                    f"recorded in {fluct['min_year']}.")
        t["s31"].append(msg)
    t["s31"].append(
        (f"The bilateral balance has consistently favoured {c['name']}, with Kenya "
         f"importing more than it exports.") if a.balance_always_negative() else "")

    # ---- Section 3.2 / 3.3 / figure notes --------------------------------
    s32_prods = top_phrase(
        a.top(a.table5, 5),
        lambda d: f"{clean_label(d['label'])} ({pct(d['share'])})")
    t["s32"] = ([f"The leading products in Kenya's exports to {c['name']} were {s32_prods}."]
                if s32_prods else [])

    s33_prods = top_phrase(
        a.top(a.table6, 5),
        lambda d: f"{clean_label(d['label'])} ({pct(d['share'])})")
    t["s33"] = ([f"Kenya's principal imports from {c['name']} were {s33_prods}."]
                if s33_prods else [])

    # ---- Section 3.3a: Goods export/import concentration & diversity -----
    conc = []
    hx = a.kenya_exports_concentration()
    hi = a.kenya_imports_concentration()
    if hx is not None:
        effx = 1.0 / hx if hx else None
        concer = "highly concentrated" if hx >= 0.25 else \
                 ("moderately concentrated" if hx >= 0.15 else "well diversified")
        conc.append(
            f"On the trade basket with {narrative_ref(c['name'])}, the HHI of Kenya's "
            f"export products was {hx:.3f}, a {concer} structure; the effective number "
            f"of distinct product lines was {effx:.0f}." if effx else
            f"On the trade basket with {narrative_ref(c['name'])}, the HHI of Kenya's "
            f"export products was {hx:.3f}, a {concer} structure.")
    if hi is not None:
        effi = 1.0 / hi if hi else None
        conceri = "highly concentrated" if hi >= 0.25 else \
                  ("moderately concentrated" if hi >= 0.15 else "well diversified")
        conc.append(
            f"For imports from {narrative_ref(c['name'])}, the product HHI was {hi:.3f} "
            f"({conceri}), implying an effective import basket of about {effi:.0f} "
            f"products." if effi else
            f"For imports from {narrative_ref(c['name'])}, the product HHI was {hi:.3f} "
            f"({conceri}).")
    if hx is not None and hi is not None:
        conc.append(
            ("Kenya's export basket is more concentrated than its import basket, "
             "indicating greater dependence on a few export lines."
             if hx > hi else
             "Kenya's import basket is the more concentrated of the two, indicating "
             "less symmetric sourcing across products."))
    # Coverage of the partner's top imported products already supplied by Kenya
    kimp_all = a.top_import_products(n=20)
    covered = [r for r in kimp_all if r.get("kenya_share") is not None and r["kenya_share"] > 0]
    if covered:
        conc.append(
            f"Kenya already supplies {len(covered)} of the partner's top 20 imported "
            f"products, showing existing lines to build on for deeper market penetration.")
    t["s33_conc"] = conc

    _t10x = a.top10_export_share()
    _hhix = a.kenya_exports_concentration()
    _cagrx = a.exports_product_cagr()
    note_x = []
    if _t10x:
        note_x.append(
            f"Top 10 export products accounted for {pct(_t10x)} "
            f"of Kenya's total exports to {narrative_ref(c['name'])} in {Y}, "
            f"underlining a {('concentrated' if _t10x >= 0.70 else 'fairly diversified')} "
            f"export basket.")
    if _hhix is not None:
        note_x.append(
            f"Measured across Kenya's bilateral export products, the Herfindahl-Hirschman "
            f"index (HHI) stood at {_hhix:.3f}, "
            f"indicating {'high' if _hhix >= 0.25 else ('moderate' if _hhix >= 0.15 else 'low')} "
            f"concentration.")
    if _cagrx is not None:
        note_x.append(
            f"Kenya's bilateral exports to {narrative_ref(c['name'])} recorded a "
            f"{_cagrx*100:,.1f}% compound annual growth rate over the series.")
    t["fig2_note"] = " ".join(note_x) if note_x else ""

    _t10i = a.top10_import_share()
    _hhii = a.kenya_imports_concentration()
    _cagri = a.imports_product_cagr()
    note_i = []
    if _t10i:
        note_i.append(
            f"Top 10 import products accounted for {pct(_t10i)} "
            f"of Kenya's total imports from {narrative_ref(c['name'])} in {Y}, "
            f"reflecting a {('concentrated' if _t10i >= 0.70 else 'fairly diversified')} "
            f"import mix.")
    if _hhii is not None:
        note_i.append(
            f"The Herfindahl-Hirschman index (HHI) of Kenya's bilateral import products was "
            f"{_hhii:.3f}, indicating {'high' if _hhii >= 0.25 else ('moderate' if _hhii >= 0.15 else 'low')} "
            f"concentration.")
    if _cagri is not None:
        note_i.append(
            f"Kenya's bilateral imports from {narrative_ref(c['name'])} recorded a "
            f"{_cagri*100:,.1f}% compound annual growth rate over the series.")
    t["fig3_note"] = " ".join(note_i) if note_i else ""

    return t


# ---------------------------------------------------------------------------
# Chart generation (matplotlib)
# ---------------------------------------------------------------------------
def make_chart_balance(a: Analysis, out_path):
    """Clustered column chart of Exports / Imports / Balance of Trade."""
    years = a.balance["years"]
    exports = a.balance["exports"]
    imports = a.balance["imports"]
    balance = a.balance["balance"]

    plt.rcParams["font.family"] = chart_font()
    fig, ax = plt.subplots(figsize=(7.6, 4.4), dpi=160)
    x = list(range(len(years)))
    w = 0.26
    cols = [lighten(THEME_ACCENTS[0], 0.30), lighten(THEME_ACCENTS[1], 0.30), lighten(THEME_ACCENTS[2], 0.30)]
    ax.bar([i - w for i in x], exports, width=w, label="Exports", color=cols[0])
    ax.bar([i for i in x], imports, width=w, label="Imports", color=cols[1])
    ax.bar([i + w for i in x], balance, width=w, label="Balance of Trade", color=cols[2])
    ax.set_xticks(x)
    ax.set_xticklabels([str(y) for y in years], fontsize=9)
    ax.set_ylabel("Value in USD Million", fontsize=10)
    ax.set_title(f"Kenya – {a.country} Balance of Trade (USD Million)", fontsize=12, weight="bold")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:,.1f}"))
    ax.yaxis.grid(True, linestyle="--", alpha=0.35)
    ax.set_axisbelow(True)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.10), ncol=3, frameon=False, fontsize=10)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)


def make_chart_share(a: Analysis, out_path, direction):
    """Styled doughnut chart of top product shares for exports or imports.

    Exports use an exploded doughnut (all slices offset, largest called
    out); imports a plain doughnut. Small slices are consolidated into an
    "Other products" slice and percentage callouts are suppressed on slices
    too small to hold them so labels never overlap.
    """
    table = a.table5 if direction == "exports" else a.table6
    items = table["items"][:10]
    labels = [short_product_name(d.get("label") or d.get("name"), d.get("code"),
                                 maxlen=42) for d in items]
    shares = series_shares(items)

    palette = list(THEME_ACCENTS)
    palette += [lighten(c, 0.45) for c in THEME_ACCENTS]

    plt.rcParams["font.family"] = chart_font()
    fig, ax = new_fig()
    labels, shares, wedges = draw_share_pie(
        ax, labels, shares, palette,
        style="3d_exploded", other_label="Other products", pct_inside=False)
    slice_callouts(fig, wedges, labels, shares)
    title = f"EXPORT SHARE {a.year}" if direction == "exports" else f"IMPORT SHARE {a.year}"
    ax.set_title(title, fontsize=12, weight="bold")
    finish(fig, out_path)


def pie_share_rows(items, n=10):
    """Labels + shares for the top ``n`` products, for editable pie charts.

    Mirrors the Excel deliverable's pie data so the Word pie (Figure 2/3)
    and the Excel deliverable's pie reflect the same numbers.
    """
    out = []
    for d in (items or [])[:n]:
        if d.get("share") is None:
            continue
        lab = short_product_name(d.get("label") or d.get("name"),
                                 d.get("code"), maxlen=42)
        out.append((lab, round(d["share"], 6)))
    return out


# ---------------------------------------------------------------------------
# Word document building
# ---------------------------------------------------------------------------
class ReportBuilder:
    def __init__(self, cfg, narratives):
        self.cfg = cfg
        self.n = narratives
        self.c = cfg["country"]
        self.a = None
        self.doc = Document()
        section = self.doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        self._add_page_border(section)
        self._setup_styles()

    def _add_page_border(self, section):
        """Black outside border frame ('wall') around every page.

        Uses ``w:pgBorders`` (page offset, 1.5pt single black line) so the
        frame sits at the page edge, around the header/footer too.
        """
        sect_pr = section._sectPr
        if sect_pr.find(qn("w:pgBorders")) is not None:
            return
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement("w:" + edge)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "12")
            el.set(qn("w:space"), "24")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        # schema order: pgBorders just before lnNumType / pgNumType / cols / docGrid
        anchor = None
        for tag in ("w:lnNumType", "w:pgNumType", "w:cols", "w:docGrid"):
            found = sect_pr.find(qn(tag))
            if found is not None:
                anchor = found
                break
        if anchor is not None:
            anchor.addprevious(borders)
        else:
            sect_pr.append(borders)

    # -- styling ------------------------------------------------------------
    def _setup_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = "Century Gothic"
        normal.font.size = Pt(12)
        rpr = normal.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rf.set(qn(attr), "Century Gothic")
        pf = normal.paragraph_format
        pf.space_after = Pt(8)
        pf.line_spacing = 1.16

        h1 = self.doc.styles["Heading 1"]
        h1.font.name = "Century Gothic"
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)
        h1.paragraph_format.keep_with_next = True

        h2 = self.doc.styles["Heading 2"]
        h2.font.name = "Century Gothic"
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(4)
        h2.paragraph_format.keep_with_next = True

        # Caption styles feed the automated List of Tables / List of Figures
        # (TOC \t fields match paragraphs by style name).
        for style_name in ("Caption Table", "Caption Figure"):
            st = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = self.doc.styles["Normal"]
            st.font.name = "Century Gothic"
            st.font.size = Pt(12)
            st.font.italic = True
            st.paragraph_format.space_before = Pt(10)
            st.quick_style = True

    def _style_run(self, run, size=None, bold=None, italic=None, color=None, name="Century Gothic"):
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color

    def add_letterhead(self):
        """KEPROBA letterhead in the running header of every page.

        Replicates the template report: the letterhead image inline in the
        default page header, 2.76in x 0.72in, left-aligned. Skipped silently
        when the asset is not present.
        """
        letterhead = os.path.join(BASE_DIR, "keproba_letterhead.png")
        if not os.path.exists(letterhead):
            return
        header = self.doc.sections[0].header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(letterhead, width=Inches(2.76))

    def add_para(self, text="", size=None, bold=None, italic=None, color=None,
                 align=None, style=None, space_after=None, space_before=None):
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        if space_before is not None:
            p.paragraph_format.space_before = Pt(space_before)
        if text:
            r = p.add_run(text)
            self._style_run(r, size=size, bold=bold, italic=italic, color=color)
        return p

    def add_background_para(self, segments):
        """Render one background paragraph from (text, needs_research) pairs.
        Segments that still need research are highlighted in yellow so they
        stand out for manual editing."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        for text, flagged in segments:
            if not text:
                continue
            r = p.add_run(text)
            self._style_run(r)
            if flagged:
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                r.font.color.rgb = RGBColor(0x9A, 0x1F, 0x1F)
                r.font.italic = True
        return p

    def add_heading(self, text, level=1, red=False):
        p = self.doc.add_paragraph(style="Heading 1" if level == 1 else "Heading 2")
        r = p.add_run(text)
        self._style_run(r, name="Century Gothic")
        return p

    def add_bullet(self, text):
        if not text:
            return None
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        self._style_run(r)
        return p

    def add_table_caption(self, text):
        p = self.doc.add_paragraph()
        r = p.add_run(_title_case(text))
        self._style_run(r, italic=True, bold=True)
        m = re.match(r"^(Table|Figure)\s+\d+\s*[:\-–]", str(text), re.IGNORECASE)
        if m:
            # style drives the automated List of Tables / List of Figures
            p.style = self.doc.styles[
                "Caption Table" if m.group(1).lower() == "table" else "Caption Figure"]
        else:
            p.paragraph_format.space_before = Pt(10)
        return p

    def add_source(self, text=SRC):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        self._style_run(r, italic=True, size=10)
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_figure(self, image_path, width_in=6.4):
        if not image_path or not os.path.exists(image_path):
            return None
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(image_path, width=Inches(width_in))
        return p

    def add_word_chart(self, kind, title, categories, series, width_in=6.3,
                       height_in=3.6, name="Chart", colors=None,
                       hole_size=50):
        """Insert a native, editable Word chart in a centred paragraph.

        kind: "bar" / "stacked" / "line" (series = [(name, [..]), ...]) or
              "pie" / "doughnut" (series = [share, ...], categories = labels).
        colors: optional per-slice / per-series palette; ``hole_size`` only
        matters for doughnuts.
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_chart.inject_chart(p, kind, title, categories, series,
                                colors=colors, width_in=width_in,
                                height_in=height_in, name=name,
                                hole_size=hole_size)
        return p

    def page_break(self):
        self.doc.add_page_break()

    def _toc_field(self, instr, placeholder):
        """Insert a Word TOC field that populates on 'Update Field'."""
        def field_piece(fld_type=None, instr=None, text=None):
            pr = self.doc.add_paragraph()
            pr.paragraph_format.space_after = Pt(0)
            run = pr.add_run()
            if fld_type:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), fld_type)
                run._r.append(el)
            if instr:
                el = OxmlElement("w:instrText")
                el.set(qn("xml:space"), "preserve")
                el.text = instr
                run._r.append(el)
            if text:
                el = OxmlElement("w:t")
                el.text = text
                run._r.append(el)
            return pr

        field_piece(fld_type="begin")
        field_piece(instr=instr)
        field_piece(fld_type="separate")
        field_piece(text=placeholder)
        field_piece(fld_type="end")

    def add_toc(self):
        p = self.doc.add_paragraph()
        r = p.add_run("Table of Contents")
        self._style_run(r, bold=True, size=14, color=RGBColor(0x15, 0x60, 0x82))
        p.paragraph_format.space_after = Pt(12)
        self._toc_field(r'TOC \o "1-2" \h \z \u',
                        "Right-click and select 'Update Field' to generate "
                        "the Table of Contents.")

    def add_list_placeholder(self, title, kind):
        """Front-matter page: 'List of Tables' / 'List of Figures'.

        A real Word TOC field scoped to the caption style, so right-click ->
        'Update Field' fills in every entry with its exact page number.
        """
        p = self.doc.add_paragraph()
        r = p.add_run(title)
        self._style_run(r, bold=True, size=14, color=RGBColor(0x15, 0x60, 0x82))
        p.paragraph_format.space_after = Pt(12)
        if kind == "table":
            instr = r'TOC \h \z \t "Caption Table,1"'
        else:
            instr = r'TOC \h \z \t "Caption Figure,1"'
        self._toc_field(instr,
                        f"Right-click and select 'Update Field' to generate "
                        f"the {title}.")

    def start_body_section(self):
        """Begin the report body as a new document section.

        Front matter (title page, TOC, lists) sits in the first section with
        an empty footer; from here on the footer carries a PAGE field whose
        numbering restarts at 1.
        """
        sec = self.doc.add_section(WD_SECTION_START.NEW_PAGE)
        sect_pr = sec._sectPr
        pg_num = OxmlElement("w:pgNumType")
        pg_num.set(qn("w:start"), "1")
        # schema order: pgNumType belongs just before w:cols / w:docGrid
        anchor = None
        for tag in ("w:cols", "w:docGrid"):
            found = sect_pr.find(qn(tag))
            if found is not None:
                anchor = found
                break
        if anchor is not None:
            anchor.addprevious(pg_num)
        else:
            sect_pr.append(pg_num)
        self._add_page_border(sec)
        return sec

    def add_footer(self):
        """Footer of the last (body) section: directorate + page number."""
        section = self.doc.sections[-1]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = ""
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run("Research and Innovation Directorate (RI)")
        self._style_run(r, size=10, italic=True)
        r2 = p.add_run("\t")
        self._style_run(r2, size=10)
        run = p.add_run()
        el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), "begin"); run._r.append(el)
        run = p.add_run()
        el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = "PAGE"; run._r.append(el)
        run = p.add_run()
        el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), "end"); run._r.append(el)

    # -- tables -------------------------------------------------------------
    @staticmethod
    def _set_table_widths(table, widths):
        """Set fixed column widths (twips) as in the original template."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "dxa")
        tblW.set(qn("w:w"), str(sum(widths)))
        grid = tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
                gc.set(qn("w:w"), str(w))
        for row in table.rows:
            seen = set()
            for ci, cell in enumerate(row.cells):
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                cell.width = Emu(widths[ci] * 635)

    @staticmethod
    def _cell_text(cell, text, bold=False, italic=False, color=None, size=10, align=None, wrap=False):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.word_wrap = wrap
        if align is not None:
            p.alignment = align
        parts = str(text).split("\n")
        for i, part in enumerate(parts):
            if i:
                run = p.add_run()
                run.add_break()
            r = p.add_run(part)
            r.font.name = "Century Gothic"
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.italic = italic
            if color is not None:
                r.font.color.rgb = color

    def _vmerge_column(self, table, col, nrows):
        """Vertically merge one column across the top ``nrows`` rows using
        explicit ``w:vMerge`` elements.

        ``Cell.merge()`` is unreliable for top-down vertical merges in
        python-docx: it can duplicate the top cell's text on the merged cell
        and fail to emit ``vMerge continue`` on the rows below.  This helper
        keeps the text in the top row, empties the continuation rows and marks
        the run correctly so the label header spans the whole header band
        (no empty box below Rank / label / Share).
        """
        tcs = [row.cells[col]._tc for row in table.rows[:nrows]]
        for tc in tcs[1:]:
            for child in list(tc):
                tc.remove(child)
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(OxmlElement("w:vMerge"))
        topPr = tcs[0].get_or_add_tcPr()
        vm = OxmlElement("w:vMerge")
        vm.set(qn("w:val"), "restart")
        topPr.append(vm)

    # -- page fitting -------------------------------------------------------
    # Vertical budget, in inches, reserved on the table's page for the
    # section heading, the italic caption and the "Source:" line so the whole
    # table always fits on a single page with those lines.
    FIT_HEADING_IN = 0.45
    FIT_CAPTION_IN = 0.35
    FIT_SOURCE_IN = 0.32
    FIT_SLACK_IN = 0.18
    LINE_FACTOR = 1.4          # line box height in units of font size (1.16 line spacing)
    ROW_EXTRA_PT = 5.0         # para spacing + cell margins + borders per row
    SIDE_MARGIN_IN = 0.14      # total horizontal cell padding subtracted
    MIN_FONT_PT = 6.0
    MAX_FONT_PT = 10.0
    MAX_LABEL_LINES = 2        # long product labels wrap at most this much

    def _usable_height_in(self):
        s = self.doc.sections[0]
        return s.page_height.inches - s.top_margin.inches - s.bottom_margin.inches

    def _cell_snapshot(self, table):
        """List (text, width_inches) per row, de-duplicating merged cells."""
        rows = []
        for row in table.rows:
            cells = []
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                txt = "\n".join(p.text for p in cell.paragraphs)
                w = cell.width.inches if cell.width else 1.0
                cells.append((txt, w))
            rows.append(cells)
        return rows

    def _lines_needed(self, text, width_in, size_pt):
        eff = max(0.4, width_in - self.SIDE_MARGIN_IN)
        cpl = max(1, int(eff * 144.0 / size_pt))
        total = 0
        for part in text.split("\n"):
            total += max(1, (len(part) + cpl - 1) // cpl)
        return max(1, total)

    def _row_height_in(self, cells, size_pt):
        lines = 1
        for txt, w in cells:
            lines = max(lines, self._lines_needed(txt, w, size_pt))
        return (lines * size_pt * self.LINE_FACTOR + self.ROW_EXTRA_PT) / 72.0

    def _table_height_in(self, rows, size_pt):
        return sum(self._row_height_in(cells, size_pt) for cells in rows)

    def _fit_font_size(self, rows, available_in):
        size = self.MAX_FONT_PT
        while size >= self.MIN_FONT_PT - 1e-9:
            if self._table_height_in(rows, size) <= available_in:
                return size
            size -= 0.5
        return self.MIN_FONT_PT

    def _truncate_overlong_cells(self, table, size_pt):
        """Shorten long product labels so no cell wraps past MAX_LABEL_LINES.

        Only invoked when a table cannot fit even at the minimum font; keeps
        the first run's formatting and cuts at a whole word -- no trailing
        ellipsis.
        """
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                w = cell.width.inches if cell.width else 1.0
                p = cell.paragraphs[0]
                text = "".join(r.text for r in p.runs)
                if not text.strip():
                    continue
                if self._lines_needed(text, w, size_pt) <= self.MAX_LABEL_LINES:
                    continue
                eff = max(0.4, w - self.SIDE_MARGIN_IN)
                cpl = max(1, int(eff * 144.0 / size_pt))
                keep = max(1, cpl * self.MAX_LABEL_LINES - 1)
                cut = text[:keep]
                sp = cut.rfind(" ")
                if sp > int(keep * 0.4):
                    cut = cut[:sp]
                cut = cut.rstrip(" ,;")
                first = p.runs[0] if p.runs else None
                for r in list(p.runs):
                    r._r.getparent().remove(r._r)
                r = p.add_run(cut)
                if first is not None:
                    r.font.name = first.font.name or "Century Gothic"
                    r.font.bold = first.font.bold
                    r.font.italic = first.font.italic
                    if first.font.color is not None and first.font.color.rgb is not None:
                        r.font.color.rgb = first.font.color.rgb
                r.font.size = Pt(size_pt)

    def _apply_fit(self, table, rows, size_pt):
        """Write the chosen row heights, font sizes and cell margins."""
        for row, cells in zip(table.rows, rows):
            h_in = self._row_height_in(cells, size_pt)
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            for tag in ("w:trHeight", "w:cantSplit"):
                for el in trPr.findall(qn(tag)):
                    trPr.remove(el)
            trPr.append(OxmlElement("w:cantSplit"))
            trH = OxmlElement("w:trHeight")
            trH.set(qn("w:val"), str(int(round(h_in * 1440))))
            trH.set(qn("w:hRule"), "atLeast")
            trPr.append(trH)

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    for r in p.runs:
                        r.font.size = Pt(size_pt)

        tblPr = table._tbl.tblPr
        for el in tblPr.findall(qn("w:tblCellMar")):
            tblPr.remove(el)
        mar = OxmlElement("w:tblCellMar")
        for tag, val in (("w:top", 12), ("w:bottom", 12),
                         ("w:left", 56), ("w:right", 56)):
            el = OxmlElement(tag)
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        look = tblPr.find(qn("w:tblLook"))
        if look is not None:
            look.addprevious(mar)
        else:
            tblPr.append(mar)

    def _fit_table_on_page(self, table):
        """Shrink the font (and, if unavoidable, long labels) so the table
        fits one page, leaving room for the heading, caption and source line.
        Each row is prevented from splitting across pages.
        """
        rows = self._cell_snapshot(table)
        available = (self._usable_height_in()
                     - self.FIT_HEADING_IN - self.FIT_CAPTION_IN
                     - self.FIT_SOURCE_IN - self.FIT_SLACK_IN)
        size = self._fit_font_size(rows, available)
        if self._table_height_in(rows, size) > available:
            self._truncate_overlong_cells(table, self.MIN_FONT_PT)
            rows = self._cell_snapshot(table)
            size = self._fit_font_size(rows, available)
        self._apply_fit(table, rows, size)

    def add_market_table(self, a: Analysis, parsed, importer=True, n_title=None, unit_row=False):
        """Table 1 (import source markets) / Table 3 (export destinations)."""
        c = self.c
        years = parsed["years"] or YEAR_HEADER
        n = len(years)
        last = years[-1]
        ncols = n + 3
        rows = 2 + int(unit_row) + len(parsed["data"])   # 2-3 header rows + data
        table = self.doc.add_table(rows=rows, cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        if importer:
            verb = "List of Exporters to"
            label_col = "Exporters"
            widths = [683, 2411] + [1077, 1077, 1078, 1078, 1080, 1080,
                                    1080, 1080, 1080, 1080][:n] + [866]
            kenya_color = BRIGHT_RED
        else:
            verb = "Importing Countries from"
            label_col = "Importers"
            widths = [730, 2430] + [1050, 1050, 1053, 1053, 1053, 1053,
                                    1053, 1053, 1053, 1053][:n] + [931]
            kenya_color = RED
        self._set_table_widths(table, widths)
        title_n = n_title or a.country_top_markets_title_n(parsed)

        # header row 0
        self._cell_text(table.cell(0, 0), f"Rank in {last}", bold=True)
        self._cell_text(table.cell(0, 1), label_col, bold=True, wrap=True)
        hdr = table.cell(0, 2)
        title = f"{verb} {c['name']}"
        if not unit_row:
            title += "\nValue in USD Billion"
        self._cell_text(hdr, title, bold=True)
        hdr.merge(table.cell(0, 1 + n))
        sh = table.cell(0, 2 + n)
        self._cell_text(sh, f"Share in {last} %", bold=True)

        # header row 1 (years)
        for k in range(n):
            self._cell_text(table.cell(1, 2 + k), str(years[k]), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        if unit_row:
            # header row 2 (unit label), as in the original template
            unit = table.cell(2, 2)
            self._cell_text(unit, "Value in USD Billion", bold=True)
            unit.merge(table.cell(2, 1 + n))
            self._cell_text(table.cell(2, 2 + n), "%", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        # Vertically merge the label + share columns across the whole header
        # band so no empty cell appears below "Rank in ...", "Exporters"/
        # "Importers" or "Share in ...", matching the template.
        last_hdr = 2 if unit_row else 1
        for c in (0, 1, 2 + n):
            self._vmerge_column(table, c, last_hdr + 1)

        # data rows
        for ri, d in enumerate(parsed["data"]):
            r = 3 + ri if unit_row else 2 + ri
            is_kenya = d["name"].lower() == "kenya"
            red = kenya_color if is_kenya else None
            bold = is_kenya
            self._cell_text(table.cell(r, 0), str(d["rank"]) if d["rank"] is not None else "",
                            bold=bold, color=red, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 1), d["name"], bold=bold, color=red, wrap=True)
            for k in range(n):
                self._cell_text(table.cell(r, 2 + k), num(d["years"][k]),
                                bold=bold, color=red, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 2 + n), pct(d["share"]), bold=True, color=red,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        self._fit_table_on_page(table)
        return table

    def add_africa_focus_table(self, a: Analysis, table=None):
        """Compact Table showing the partner's exports destined to Africa,
        with the Kenya row highlighted.  Backs the 'Africa / Kenya' insight."""
        table = table or a.table3
        afr = a.africa_destinations(table)
        kenya = a.kenya_destination(table)
        if kenya and kenya not in afr:
            afr = afr + [kenya]
        # order by rank, Kenya last labelled
        afr = [d for d in afr if d.get("share") is not None]
        afr = sorted(afr, key=lambda d: (d["name"].lower() != "kenya", d.get("rank") or 0))
        if not afr:
            return None
        n = len(afr) + 2                      # header + data rows + total row
        t = self.doc.add_table(rows=n, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_widths(t, [730, 3600, 2440])
        self._cell_text(t.cell(0, 0), f"Rank in {a.year}", bold=True)
        self._cell_text(t.cell(0, 1), "African destination", bold=True)
        self._cell_text(t.cell(0, 2), f"Share of {a.country} exports in {a.year} %", bold=True)
        afr_total = a.africa_export_share(table)
        for ri, d in enumerate(afr, start=1):
            is_kenya = d["name"].lower() == "kenya"
            red = RED if is_kenya else None
            self._cell_text(t.cell(ri, 0), str(d.get("rank") or ""),
                            bold=is_kenya, color=red, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(t.cell(ri, 1), display_name(d["name"]), bold=is_kenya,
                            color=red, wrap=True)
            self._cell_text(t.cell(ri, 2), pct(d["share"]), bold=is_kenya, color=red,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        rsum = len(afr) + 1
        self._cell_text(t.cell(rsum, 0), "", bold=True)
        self._cell_text(t.cell(rsum, 1), "Total, Africa", bold=True)
        self._cell_text(t.cell(rsum, 2), pct(afr_total), bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        self._fit_table_on_page(t)
        return t

    def add_product_table(self, a: Analysis, parsed, flow_label, unit="Value in USD Billion",
                          unit_row=False, widths=None):
        """Tables 2, 4, 5, 6 (product tables)."""
        years = parsed["years"] or YEAR_HEADER
        n = len(years)
        last = years[-1]
        ncols = n + 4
        rows = 2 + int(unit_row) + len(parsed["data"])
        table = self.doc.add_table(rows=rows, cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if widths is None:
            widths = [685, 972, 3170, 731, 731, 731, 731, 733,
                      733, 733, 733, 733, 733, 866]
        self._set_table_widths(table, fit_widths(widths, n))

        self._cell_text(table.cell(0, 0), f"Rank in {last}", bold=True)
        self._cell_text(table.cell(0, 1), "Code", bold=True)
        self._cell_text(table.cell(0, 2), "Product label", bold=True, wrap=True)
        hdr = table.cell(0, 3)
        title = flow_label if unit_row else f"{flow_label}  {unit}"
        self._cell_text(hdr, title, bold=True)
        hdr.merge(table.cell(0, 2 + n))
        sh = table.cell(0, 3 + n)
        self._cell_text(sh, f"Share in {last} %", bold=True)

        for k in range(n):
            self._cell_text(table.cell(1, 3 + k), str(years[k]), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        if unit_row:
            unit = table.cell(2, 3)
            self._cell_text(unit, "Value in USD Billion", bold=True)
            unit.merge(table.cell(2, 2 + n))
            self._cell_text(table.cell(2, 3 + n), "%", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        # Vertically merge the label + share columns across the whole header
        # band so no empty cell sits below "Rank in ...", "Code",
        # "Product label" or "Share in ...", matching the template.
        last_hdr = 2 if unit_row else 1
        for c in (0, 1, 2, 3 + n):
            self._vmerge_column(table, c, last_hdr + 1)

        for ri, d in enumerate(parsed["data"]):
            r = 3 + ri if unit_row else 2 + ri
            is_total = d["kind"] == "total"
            is_all = d["kind"] == "all_other"
            b = is_total or is_all
            self._cell_text(table.cell(r, 0), str(d["rank"]) if d["rank"] is not None else "",
                            bold=b, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 1), str(d["code"] or ""), bold=b)
            self._cell_text(table.cell(r, 2), str(d["name"]), bold=b, wrap=True)
            for k in range(n):
                self._cell_text(table.cell(r, 3 + k), num(d["years"][k]),
                                bold=b, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 3 + n), pct(d["share"]), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        self._fit_table_on_page(table)
        return table

    def add_quick_facts(self, facts):
        table = self.doc.add_table(rows=len(facts), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_widths(table, [3118, 6520])
        for ri, (k, v) in enumerate(facts):
            self._cell_text(table.cell(ri, 0), k, bold=True)
            self._cell_text(table.cell(ri, 1), v)
        # column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.2)
            row.cells[1].width = Inches(4.2)
        return table

    def _excel_cell_text(self, cell):
        """Render an Excel cell for a Word table: thousands separators for
        values >= 1000, one decimal by default, percents as '27.4%', while
        index columns (RCA/HHI/score) keep the precision their Excel format
        declares.
        """
        v = cell.value
        if v is None:
            return ""
        if isinstance(v, str):
            return v
        if isinstance(v, bool):
            return str(v)
        if isinstance(v, (int, float)):
            fmt = (cell.number_format or "").replace("General", "").strip()
            if "%" in fmt:
                return pct(v, 1)
            dec = 1
            m = re.search(r"0+(\.0+)?", fmt)
            if m and m.group(1):
                dec = len(m.group(1).strip("."))
            elif fmt and not re.search(r"\.0+", fmt):
                dec = 0
            return num(v, dec)
        return str(v)

    def add_table_from_excel(self, excel_path, sheet_name=None):
        """Add a table from an Excel file to the document."""
        import openpyxl as xl
        wb = xl.load_workbook(excel_path, data_only=True)
        if sheet_name and sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            ws = wb.active
        rows = []
        for row in ws.iter_rows():
            rows.append([self._excel_cell_text(c) for c in row])
        wb.close()
        if not rows:
            return
        n_rows = len(rows)
        n_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci < n_cols:
                    self._cell_text(table.cell(ri, ci), val)
        return table

    def add_bilateral_table(self, a: Analysis):
        """Table 7 (Kenya's export composition & market alignment).

        Styled like Tables 1-6: bold header, one decimal for values,
        percentages as '12.3%', widths fitted to the page.
        """
        c = self.c
        t7 = a.table7
        y = t7.get("year") or self.a.year
        items = t7["items"]
        table = self.doc.add_table(rows=1 + len(items), cols=7)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_widths(
            table, [640, 900, 3100, 1560, 1130, 1230, 930])

        headers = [
            f"Rank in {y}",
            "HS Code",
            "Product label",
            f"Kenya's exports to {c['name']}\nValue in USD Million",
            f"Share of Kenya's Exports to {c['name']}",
            f"Share in {c['name']}'s Imports",
            "Annual Growth %",
        ]
        for ci, text in enumerate(headers):
            self._cell_text(table.cell(0, ci), text, bold=True, wrap=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        for ri, it in enumerate(items):
            r = 1 + ri
            self._cell_text(table.cell(r, 0), str(it["rank"]),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 1), str(it["code"]),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 2), it["label"], wrap=True)
            self._cell_text(table.cell(r, 3), num(it["value"]),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 4), pct(it["kenya_share"]),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 5), pct(it["partner_share"]),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
            self._cell_text(table.cell(r, 6), pct(it["growth"]),
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        self._fit_table_on_page(table)
        return table


def bilateral_bullets(a):
    """Interpretation bullets for Table 7 (export composition & alignment)."""
    t7 = getattr(a, "table7", None)
    if not t7 or not t7["items"]:
        return []
    y = t7.get("year") or a.year
    items = [it for it in t7["items"] if it["value"] is not None]
    if not items:
        return []
    bullets = []

    top = items[:10]
    top_val = sum(it["value"] for it in top)
    top_share = sum(it["kenya_share"] for it in top if it["kenya_share"] is not None)
    bullets.append(
        f"Kenya's top {len(top)} export products to {a.country} were worth "
        f"USD {top_val:,.1f} million in {y}, together representing "
        f"{pct(top_share)} of Kenya's exports to {a.country}.")

    aligned = [it for it in top if it["partner_share"] is not None
               and it["partner_share"] > 0]
    if aligned:
        best = max(aligned, key=lambda it: it["partner_share"])
        bullets.append(
            f"{best['label']} shows the strongest market alignment: it takes "
            f"{pct(best['kenya_share'])} of Kenya's exports to {a.country} "
            f"while accounting for {pct(best['partner_share'])} of "
            f"{a.country}'s import demand in {y}.")

    growing = [it for it in items if it["growth"] is not None]
    if growing:
        fastest = max(growing, key=lambda it: it["growth"])
        bullets.append(
            f"Fastest growing export to {a.country} was {fastest['label']} "
            f"({pct(fastest['growth'])} year-on-year), suggesting momentum "
            "worth consolidating.")
    return bullets


# ---------------------------------------------------------------------------
# Main report assembly
# ---------------------------------------------------------------------------
def build_report(cfg, excel_dir, out_path, tmp_dir):
    os.makedirs(tmp_dir, exist_ok=True)
    a = Analysis(cfg).load(excel_dir)
    narr = build_narratives(a, cfg)

    c = cfg["country"]
    rep = cfg["report"]
    Y = a.year

    b = ReportBuilder(cfg, narr)
    b.a = a
    doc = b.doc

    # ============================== TITLE PAGE ==============================
    b.add_letterhead()
    b.add_para(f"KENYA- {title_partner(c['name'])} TRADE FLOW",
               size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    b.add_para(rep["title_line2"], size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    for _ in range(3):
        b.add_para("")
    b.add_para("Prepared", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    b.add_para("BY", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    b.add_para("Kenya Export Promotion and Branding Agency (KEPROBA)",
               size=22, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    b.add_para("Research and Innovation Directorate (RID)",
               size=22, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    b.add_para(rep["month_year"], size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    b.page_break()

    # ============================== TABLE OF CONTENTS =======================
    b.add_toc()
    b.page_break()

    # ==================== LIST OF TABLES / LIST OF FIGURES ==================
    b.add_list_placeholder("List of Tables", "table")
    b.page_break()
    b.add_list_placeholder("List of Figures", "figure")
    # page numbers start at 1 from here (title/TOC/lists carry none)
    b.start_body_section()

    # ============================== SECTION 1 ===============================
    b.add_heading(f"1. {c['title']}")
    b.add_heading("1.1. Background", level=2)
    for segments in narr["background"]:
        b.add_background_para(segments)
    b.page_break()

    # ============================== SECTION 2 ===============================
    b.add_heading("2. Situational Analysis")
    b.add_para(f"{c['name']}'s Trade Flows and Position in the Global Trade", bold=True)
    b.add_para("Exports", bold=True, color=BRIGHT_RED)
    for line in narr["exports_bullets"]:
        b.add_bullet(line)
    b.add_para("Imports", bold=True, color=BRIGHT_RED)
    for line in narr["imports_bullets"]:
        b.add_bullet(line)
    b.page_break()

    # 2.1 Import source markets
    b.add_heading(f"2.1 {c['name']}'s Lead Import Source Markets", level=2)
    n = a.country_top_markets_title_n(a.table1)
    b.add_table_caption(f"Table 1: {a.possessive} top {n} Import Source Markets")
    b.add_market_table(a, a.table1, importer=True)
    b.add_source()
    for line in narr["s21"]:
        b.add_bullet(line)
    for line in narr.get("s21_kenya", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # 2.2 Import products
    b.add_heading(f"2.2 {c['name']} Lead Import products from the world", level=2)
    b.add_table_caption(f"Table 2: {a.possessive} top {len(a.table2['items'])} Import products from the World.")
    b.add_product_table(a, a.table2, "Imported Products from the World", "Value in USD Billion",
                        unit_row=True, widths=[685, 972, 3170, 731, 731, 731, 731, 733, 866])
    b.add_source()
    for line in narr["s22"]:
        b.add_bullet(line)
    for line in narr.get("s22_kenya", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # 2.3 Export destinations
    b.add_heading(f"2.3 {c['name']}'s Lead Export Destination Markets", level=2)
    n = a.country_top_markets_title_n(a.table3)
    b.add_table_caption(f"Table 3: {a.short}'s top {n} Export Destination Markets")
    b.add_market_table(a, a.table3, importer=False, unit_row=True)
    b.add_source()
    for line in narr["s23"]:
        b.add_bullet(line)
    if a.africa_destinations():
        b.add_para("Focus: {name}'s exports to Africa & Kenya".format(name=c["name"]), bold=True)
        b.add_table_caption(f"Table 3a: {a.country}'s exports destined to Africa in {Y}")
        b.add_africa_focus_table(a)
        b.add_source()
    b.page_break()

    # 2.4 Export products
    b.add_heading(f"2.4 {c['name']}'s Lead Export Products to the World", level=2)
    b.add_table_caption(f"Table 4: {a.short}'s top {len(a.table4['items'])} export products to the world")
    b.add_product_table(a, a.table4, "Exported Products to the World", "Value in USD Billion",
                        widths=[730, 1048, 2707, 711, 711, 711, 711, 711, 931])
    b.add_source()
    for line in narr["s24"]:
        b.add_bullet(line)
    b.page_break()

    # ============================== SECTION 3 ===============================
    b.add_heading(f"3. Bilateral Trade Between Kenya and {c['title']}")

    # 3.1 Trends
    b.add_heading(f"3.1 Kenya – {c['name']} Bilateral Trade Trends", level=2)
    b.add_table_caption(f"Figure 1: Kenya – {c['name']} Balance of Trade")
    b.add_word_chart(
        "bar", f"Kenya – {c['name']} Balance of Trade (USD Million)",
        a.balance["years"],
        [("Exports", a.balance["exports"]),
         ("Imports", a.balance["imports"]),
         ("Balance of Trade", a.balance["balance"])],
        width_in=6.3, height_in=3.6, name="Figure 1 - Balance of Trade")
    b.add_source()
    for line in narr["s31"]:
        b.add_bullet(line)

    # Export share figure (placed after Table 5, mirroring Table 6 -> Figure 3)
    # 3.2 Export products
    b.add_heading(f"3.2 Kenya's Export Products to {c['name']}", level=2)
    b.add_table_caption(f"Table 5: Kenya's Top {len(a.table5['items'])} Export Products to {c['name']} in {Y}")
    b.add_product_table(a, a.table5, f"Kenya's exports to {c['name']}", "Value in USD Million",
                        widths=[730, 2094, 2094, 656, 711, 711, 711, 712, 931])
    b.add_source()
    for line in narr["s32"]:
        b.add_bullet(line)

    b.add_para(f"Share of Kenya's Top Export Products to {c['name']}", bold=True)
    b.add_table_caption(f"Figure 2: Share of Kenya's Top Exports to {c['name']} in {Y}")
    _ex2 = pie_share_rows(a.table5.get("items") if a.table5 else None)
    b.add_word_chart(
        "pie", f"Share of Kenya's Top Exports to {c['name']} in {Y}",
        [lab for lab, _ in _ex2], [sh for _, sh in _ex2],
        width_in=5.6, height_in=4.4, name="Figure 2 - Top Exports")
    b.add_bullet(narr["fig2_note"])
    b.page_break()

    # 3.3 Import products
    b.add_heading(f"3.3 Kenya's Key Import Products from {c['name']}", level=2)
    b.add_table_caption(f"Table 6: Kenya's Top {len(a.table6['items'])} Import Products from {c['name']} in {Y}")
    b.add_product_table(a, a.table6, f"Kenya's imports from {c['name']}", "Value in USD Million",
                        widths=[730, 1901, 1902, 821, 821, 821, 711, 712, 931])
    b.add_source()
    for line in narr["s33"]:
        b.add_bullet(line)
    for line in narr.get("s33_conc", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # 3.4 Import share figure
    b.add_heading("3.4 Share of top Import Products from " + c["name"], level=2)
    b.add_table_caption(f"Figure 3: Share of Kenya's Top Imports from {c['name']} in {Y}")
    _ex3 = pie_share_rows(a.table6.get("items") if a.table6 else None)
    b.add_word_chart(
        "pie", f"Share of Kenya's Top Imports from {c['name']} in {Y}",
        [lab for lab, _ in _ex3], [sh for _, sh in _ex3],
        width_in=5.6, height_in=4.4, name="Figure 3 - Top Imports")
    b.add_bullet(narr["fig3_note"])
    b.page_break()

    # 3.5 Kenya's export composition and market alignment
    b.add_heading(f"3.5 Kenya's Export Composition and Market Alignment with {c['name']}", level=2)
    if getattr(a, "table7", None) and a.table7["items"]:
        b.add_table_caption(f"Table 7: Kenya's Export Products – Market Alignment with {c['name']} in {Y}")
        b.add_bilateral_table(a)
        b.add_source()
        for line in bilateral_bullets(a):
            b.add_bullet(line)
    else:
        b.add_para("[Table 7 data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
        b.add_source()
    b.add_para(
        f"This table compares Kenya's export products to {c['name']} against "
        f"{c['name']}'s overall import demand. Products where Kenya has a significant "
        "share of its own exports AND the partner has high import demand represent "
        "the strongest market alignment opportunities.",
        italic=True)
    b.page_break()

    # ============================== SECTION 4 ===============================
    b.add_heading(f"4. Kenya's Export Potential on {c['title']} Market")
    ep = cfg.get("export_potential", {}) or {}
    b.add_table_caption(f"Figure 4: Kenya's products with export potential to {c['name']}.")
    img = resolve_path(cfg_path, ep.get("image"))
    if img:
        b.add_figure(img)
    else:
        b.add_para("[Placeholder: insert ITC Export Potential chart for "
                   f"{c['name']} in config 'export_potential.image']", italic=True, color=RGBColor(0x60, 0x60, 0x60))
    b.add_source()
    for para in ep.get("paragraphs", []):
        b.add_bullet(para)
    b.page_break()

    # ============================== SECTION 5 ===============================
    b.add_heading(f"5. Annex 1: The Map of {c['name']}")
    mp = cfg.get("map", {}) or {}
    img = resolve_path(cfg_path, mp.get("image"))
    if img:
        b.add_figure(img, width_in=4.6)
    else:
        b.add_para("[Placeholder: insert map image in config 'map.image']",
                   italic=True, color=RGBColor(0x60, 0x60, 0x60))
    b.add_source(f"Source: {mp.get('source', 'Google map')}")
    b.page_break()

    # ============================== SECTION 6 ===============================
    b.add_heading("6. Annex II: Quick Facts")
    b.add_quick_facts(cfg.get("quick_facts", []))
    b.page_break()

    # ============================== REFERENCES ==============================
    b.add_heading("References")
    for ref in cfg.get("references", []):
        b.add_para(ref)

    # footer
    b.add_footer()

    doc.save(out_path)

    # Companion editable Excel workbook: tables as plain editable ranges
    # (no AutoFilter, thick outer borders) plus native editable charts.
    try:
        import excel_deliverable
        xlsx_out = os.path.splitext(out_path)[0] + " TABLES.xlsx"
        xdir = os.path.dirname(xlsx_out) or "."
        os.makedirs(xdir, exist_ok=True)
        excel_deliverable.build_deliverable(a, cfg, xlsx_out)
        print(f"      Excel deliverable saved to   : {xlsx_out}")
    except Exception as _e:  # additive; never break the report build
        print(f"      [warn] Excel deliverable skipped: {_e}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def load_config(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def main():
    ap = argparse.ArgumentParser(description="Generate a Kenya – <Country> Trade Flow report.")
    ap.add_argument("--excel-dir", default=os.path.join(BASE_DIR, "sample_data"),
                    help="Folder containing the seven ITC Excel files")
    ap.add_argument("--config", default=os.path.join(BASE_DIR, "config", "saudi_arabia.json"),
                    help="Per-country JSON configuration file")
    ap.add_argument("--output", default=None, help="Output .docx path")
    ap.add_argument("--tmp", default=os.path.join(BASE_DIR, "output", ".tmp"),
                    help="Temporary directory for chart images")
    args = ap.parse_args()

    global cfg_path
    cfg_path = os.path.abspath(args.config)
    cfg = load_config(cfg_path)

    out = args.output
    if not out:
        cname = title_partner(cfg["country"]["name"])
        out = os.path.join(BASE_DIR, "output", f"KENYA-{cname} TRADE FLOW.docx")
    out = os.path.abspath(out)
    out_dir = os.path.dirname(out)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    os.makedirs(args.tmp, exist_ok=True)

    print(f"[1/3] Reading Excel data from   : {os.path.abspath(args.excel_dir)}")
    print(f"[2/3] Building report for       : {cfg['country']['name']}")
    build_report(cfg, args.excel_dir, out, args.tmp)
    print(f"[3/3] Report saved to           : {out}")


if __name__ == "__main__":
    main()
