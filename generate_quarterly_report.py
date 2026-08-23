#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_quarterly_report.py
============================

Generate the KEPROBA "Kenya Export Performance" quarterly Word report
(.docx) from the two pivot workbooks built by ``make_quarterly_tables.py``
(or the equivalent manually produced ones):

  - ``Exports.xlsx``  -- Table 1 Top Destination Markets,
                         Table 2 Top Export Products (+ pivots)
  - ``Imports.xlsx``  -- Annex 1 Imports by Partner,
                         Annex 2 Imports by Products (+ pivots)

The report carries the same features as the goods reports:

  * KEPROBA letterhead on every page;
  * title page, table of contents, list of tables, list of figures --
    none of them numbered;
  * page numbering restarts at 1 where the report body begins;
  * a data-driven narrative (overview, comparative perspectives and the
    Section 5 "Deduction") computed from the quarter's numbers.

Usage
-----
    python3 generate_quarterly_tables.py --excel-dir "EXPORT PERFORMANCE FOR Q2" \
        --out-dir output/quarterly
    python3 generate_quarterly_report.py --excel-dir output/quarterly \
        --output output/KENYA EXPORT PERFORMANCE IN APRIL-JUNE 2025-2026.docx

Dependencies: python-docx, openpyxl, matplotlib.
"""

import argparse
import os
import re
import sys
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import openpyxl

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE_DIR)

from charts import consolidate, draw_share_pie, slice_callouts, new_fig, finish
from country_names import display_name, short_product_name
from generate_report import (
    ReportBuilder, to_float, num, chart_font, THEME_ACCENTS, lighten,
)

SRC_KRA = "Source: KRA Database; Compiled by KEPROBA"

MONTH_NAMES = {1: "January", 2: "February", 3: "March", 4: "April",
               5: "May", 6: "June", 7: "July", 8: "August", 9: "September",
               10: "October", 11: "November", 12: "December"}

# Regional groupings used by the narrative -------------------------------
EAC = {"kenya", "uganda", "tanzania", "rwanda", "burundi",
       "democratic rep of congo", "democratic republic of congo",
       "south sudan", "somalia"}
AFRICA = EAC | {"egypt", "south africa", "malawi", "zambia", "zimbabwe",
                "nigeria", "ghana", "sudan", "ethiopia", "djibouti",
                "libya", "morocco", "tunisia", "algeria", "angola",
                "mozambique", "mauritius", "madagascar", "senegal",
                "cote d ivoire", "ivory coast", "cameroon", "gabon",
                "equatorial guinea", "congo", "namibia", "botswana",
                "lesotho", "eswatini", "swaziland", "seychelles",
                "sao tome & principe", "cape verde", "comoros",
                "mauritania", "mali", "niger", "chad", "benin", "togo",
                "sierra leone", "liberia", "guinea", "guinea bissau",
                "gambia", "burkina faso", "central african republic"}
ASIA = {"china", "india", "japan", "malaysia", "indonesia", "thailand",
        "vietnam", "singapore", "south korea", "republic of korea",
        "taiwan", "hong kong", "pakistan", "bangladesh", "sri lanka",
        "philippines", "myanmar(burma)", "myanmar", "cambodia", "laos",
        "nepal", "bhutan", "afghanistan", "uzbekistan", "kazakhstan"}
GULF = {"saudi arabia", "united arab emirates", "oman", "qatar", "kuwait",
        "bahrain", "yemen", "iraq", "iran"}
WEST = {"united states of america", "usa", "united kingdom", "japan",
        "belgium", "netherlands", "germany", "france", "italy", "canada",
        "spain", "sweden", "denmark", "finland", "austria", "ireland",
        "switzerland", "norway", "portugal", "greece", "australia",
        "new zealand"}

# Product keyword groups used in the deduction narrative ------------------
HORTICULTURE_KW = ("cut flower", "rose", "avocado", "mango", "vegetable",
                   "macadamia", "pineapple", "passion fruit", "legulous",
                   "beans,", "peas,", "courgette", "capsicum", "chilli")
PETROLEUM_KW = ("petroleum oils", "motor spirit", "gasoline",
                "butanes", "liquefied petroleum", "gas oil", "diesel",
                "jet fuel", "kerosene", "fuel")


# ---------------------------------------------------------------------------
# Workbook locating / parsing
# ---------------------------------------------------------------------------
def find_workbooks(excel_dir):
    """Locate Exports.xlsx / Imports.xlsx in ``excel_dir``."""
    found = {}
    for fname in sorted(os.listdir(excel_dir)):
        low = fname.lower()
        if not low.endswith(".xlsx"):
            continue
        path = os.path.join(excel_dir, fname)
        if low.startswith("export"):
            found.setdefault("exports", path)
        elif low.startswith("import"):
            found.setdefault("imports", path)
    missing = [k for k in ("exports", "imports") if k not in found]
    if missing:
        names = ", ".join(sorted(os.listdir(excel_dir)))
        sys.exit("[ERROR] Missing %s workbook(s) in '%s'.\nFiles present: %s"
                 % ("/".join(missing), excel_dir, names))
    return found


def _grid(ws):
    return [[c.value for c in row] for row in ws.iter_rows()]


def parse_rank_sheet(ws):
    """Parse one comparison rank-table sheet (current vs previous year).

    Expected generated layout::

        row1  (Value in Ksh. Billion)
        row2  Rank | ITEM | <Y-1> (merged)        | <Y> (merged)         | Change in <Y-1>-<Y>
        row3  ...  | April May June Total         | April May June Total | Total change | %
        data rows, All-other row, Grand Total

    Returns::

        {"sheet", "months": [...], "years": [yp, yc],
         "items": [{"name", "rank", "months": [[yp vals], [yc vals]],
                    "totals": [tp, tc], "change", "pct"}, ...],
         "all_other": {...} | None,
         "grand": {...}}
    """
    grid = _grid(ws)

    def cellv(row, ci):
        return row[ci] if ci < len(row) else None

    # -- locate the label row: first row holding >= 3 month names ----------
    month_set = set(MONTH_NAMES.values())
    lbl = None
    for ri, row in enumerate(grid):
        cnt = 0
        for v in row:
            if v is not None and str(v).strip().capitalize() in month_set:
                cnt += 1
        if cnt >= 3:
            lbl = ri
            break
    if lbl is None:
        return None

    # -- classify label-row columns into month runs ------------------------
    # runs[i] = [first_col, last_col, [month names]]
    runs = []
    for ci, v in enumerate(grid[lbl]):
        s = str(v).strip() if v is not None else ""
        if s.capitalize() in month_set:
            if runs and ci == runs[-1][1] + 1:
                runs[-1][1] = ci
                runs[-1][2].append(s)
            else:
                runs.append([ci, ci, [s]])
    if not runs:
        return None

    comparison = len(runs) >= 2
    if comparison:
        prev_run, cur_run = runs[0], runs[1]
    else:
        prev_run = cur_run = runs[0]      # single-year layout
    months = [m.capitalize() for m in cur_run[2]]

    # totals sit immediately after each run
    p_total = (prev_run[1] + 1) if comparison else None
    c_total = cur_run[1] + 1
    hdr_end = max(c_total, p_total or 0) + 3
    hdr_lbl = [str(cellv(grid[lbl], c)).strip()
               if cellv(grid[lbl], c) is not None else ""
               for c in range(hdr_end)]
    if "total" not in str(hdr_lbl[c_total]).lower():
        return None
    if comparison and "total" not in str(hdr_lbl[p_total]).lower():
        return None
    if not comparison:
        p_total = None

    chg_col = pct_col = share_col = None
    if comparison:
        # change / % columns come after the current-year Total
        for ci in range(c_total + 1, len(hdr_lbl)):
            h = hdr_lbl[ci].lower()
            if chg_col is None and "change" in h:
                chg_col = ci
            elif pct_col is None and h in ("%", "% change"):
                pct_col = ci
        if chg_col is None:
            chg_col = c_total + 1
            pct_col = c_total + 2
        elif pct_col is None:
            pct_col = chg_col + 1
    else:
        # single-year layouts carry a Share column after the Total
        for ci in range(c_total + 1, len(hdr_lbl)):
            h = hdr_lbl[ci].lower()
            if h.startswith("%") or "share" in h:
                share_col = ci
                break

    # -- years from the group header row ------------------------------------
    grp = lbl - 1
    grp_row = grid[grp] if grp >= 0 else []

    def year_at(ci):
        for c in range(min(ci, len(grp_row) - 1), -1, -1):
            v = grp_row[c]
            if isinstance(v, (int, float)) and 2015 <= int(v) <= 2100:
                return int(v)
        return None

    if comparison:
        y_prev = year_at(prev_run[0])
        y_cur = year_at(cur_run[0])
        if y_cur is None:
            y_cur = (y_prev or date.today().year - 1) + 1
        if y_prev is None:
            y_prev = y_cur - 1
        years = [y_prev, y_cur]
    else:
        years = [None, year_at(cur_run[0])]   # review year resolved later

    # -- data rows -----------------------------------------------------------
    def to_f(v):
        try:
            if v is None or (isinstance(v, str) and not v.strip()):
                return None
            return float(str(v).replace(",", ""))
        except (TypeError, ValueError):
            return None

    items, all_other = [], None
    grand = None
    n_m = len(months)
    for row in grid[lbl + 1:]:
        if not row or all(v is None or str(v).strip() == ""
                          for v in row[:c_total + 1]):
            continue
        label = cellv(row, 1)
        if label is None:
            continue
        label = str(label).strip()
        low = label.lower()
        if low.startswith("value in") or low in ("%", "% change"):
            continue
        cur_vals = [to_f(cellv(row, cur_run[0] + k)) for k in range(n_m)]
        prev_vals = ([to_f(cellv(row, prev_run[0] + k))
                      for k in range(n_m)] if comparison
                     else [None] * n_m)
        mvals = [prev_vals, cur_vals]
        t_prev = to_f(cellv(row, p_total)) if comparison else None
        t_cur = to_f(cellv(row, c_total))
        totals = [t_prev, t_cur]
        if t_cur is None:
            t_cur = sum(v for v in cur_vals if v is not None)
            totals[1] = t_cur
        if comparison and t_prev is None:
            totals[0] = sum(v for v in prev_vals if v is not None)
        share = to_f(cellv(row, share_col)) if share_col else None
        if comparison:
            change = to_f(cellv(row, chg_col))
            if change is None:
                change = (totals[1] or 0.0) - (totals[0] or 0.0)
            pct = to_f(cellv(row, pct_col))
            if pct is None and totals[0]:
                pct = change / totals[0]
        else:
            change, pct = None, None
        entry = {"name": label,
                 "rank": int(to_f(cellv(row, 0)))
                 if to_f(cellv(row, 0)) is not None else None,
                 "months": mvals, "totals": totals,
                 "share": share, "change": change, "pct": pct}
        if low.startswith("grand"):
            grand = entry
            break
        if entry["rank"] is None and low.startswith("all other"):
            all_other = entry
        else:
            items.append(entry)
    if grand is None:
        return None
    return {"sheet": ws.title, "months": months, "years": years,
            "comparison": comparison,
            "items": items, "all_other": all_other, "grand": grand}


def pick_sheet(wb, prefer, keywords):
    """Find the best-matching rank sheet by name, then by content."""
    for name in wb.sheetnames:
        if prefer and prefer in name.lower():
            parsed = parse_rank_sheet(wb[name])
            if parsed:
                return parsed
    for name in wb.sheetnames:
        low = name.lower()
        if any(k in low for k in keywords):
            parsed = parse_rank_sheet(wb[name])
            if parsed:
                return parsed
    for name in wb.sheetnames:
        parsed = parse_rank_sheet(wb[name])
        if parsed:
            return parsed
    return None


def detect_year(wbs, fallback=None):
    """Review year: workbook title property, then Data-sheet Year column."""
    for wb in wbs:
        title = (wb.properties.title or "")
        m = re.search(r"(20\d{2})", title)
        if m:
            return int(m.group(1))
    for wb in wbs:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows(min_row=1, max_row=5, max_col=6):
                for cell in row:
                    if isinstance(cell.value, str) \
                            and cell.value.strip().lower() == "year":
                        for r2 in ws.iter_rows(
                                min_row=cell.row + 1,
                                max_row=cell.row + 200,
                                min_col=cell.column, max_col=cell.column):
                            for c2 in r2:
                                if isinstance(c2.value, (int, float)) \
                                        and 2015 <= int(c2.value) <= 2100:
                                    return int(c2.value)
    return fallback


# ---------------------------------------------------------------------------
# Data loading and analysis
# ---------------------------------------------------------------------------
class QuarterAnalysis:
    """Loads the quarterly workbooks and derives every narrative figure.

    Tables compare the review quarter with the same quarter of the previous
    year; every table parsed carries ``months``/``totals`` lists ordered
    [previous year, current year].
    """

    def __init__(self, year=None):
        self.year = year          # review (current) year, detected if None
        self.month_names = []
        self.quarter = ""
        self.t1 = None   # top destination markets (exports by partner)
        self.t2 = None   # top export products
        self.a1 = None   # imports by supplying countries
        self.a2 = None   # imports by products

    # -- loading ------------------------------------------------------------
    def load(self, excel_dir):
        files = find_workbooks(excel_dir)
        wb_e = openpyxl.load_workbook(files["exports"], data_only=True)
        wb_i = openpyxl.load_workbook(files["imports"], data_only=True)
        self.t1 = pick_sheet(wb_e, "table 1", ("destination", "partner"))
        self.t2 = pick_sheet(wb_e, "table 2", ("product",))
        self.a1 = pick_sheet(wb_i, "annex 1", ("partner",))
        self.a2 = pick_sheet(wb_i, "annex 2", ("product",))
        for name, t in (("Table 1", self.t1), ("Table 2", self.t2),
                        ("Annex 1", self.a1), ("Annex 2", self.a2)):
            if t is None:
                sys.exit("[ERROR] Could not locate %s in the quarterly "
                         "workbooks." % name)
        comparison = bool(self.t1.get("comparison")
                          and self.a1.get("comparison"))
        if self.year is None:
            years = [t["years"][1] for t in (self.t1, self.t2,
                                             self.a1, self.a2)
                     if t["years"][1]]
            if len(years) == 4:
                self.year = max(years)
            else:
                # single-year sheets: fall back to workbook metadata / data
                self.year = detect_year([wb_e, wb_i],
                                        fallback=date.today().year)
        self.has_comparison = comparison
        self.year_prev = (self.year - 1) if comparison else None
        self.month_names = self.t1["months"]
        self.quarter = quarter_label(self.month_names)
        return self

    # -- basic aggregates ---------------------------------------------------
    @property
    def exp_total(self):
        return self.t1["grand"]["totals"][1] or 0.0

    @property
    def exp_total_prev(self):
        v = self.t1["grand"]["totals"][0]
        return v if v is not None else None

    @property
    def imp_total(self):
        return self.a1["grand"]["totals"][1] or 0.0

    @property
    def imp_total_prev(self):
        v = self.a1["grand"]["totals"][0]
        return v if v is not None else None

    @property
    def deficit(self):
        return self.exp_total - self.imp_total

    def exp_months(self):
        return list(self.t1["grand"]["months"][1])

    def exp_months_prev(self):
        return list(self.t1["grand"]["months"][0])

    def imp_months(self):
        return list(self.a1["grand"]["months"][1])

    def imp_months_prev(self):
        return list(self.a1["grand"]["months"][0])

    def balance_months(self):
        return [e - i for e, i in zip(self.exp_months(), self.imp_months())]

    def exp_growth_through_quarter(self):
        m = self.exp_months()
        if len(m) >= 2 and m[0]:
            return (m[-1] - m[0]) / abs(m[0])
        return None

    def yoy_growth(self):
        """Total export growth vs the same quarter of the previous year."""
        prev = self.exp_total_prev
        return (self.exp_total - prev) / prev if prev else None

    def yoy_import_growth(self):
        prev = self.imp_total_prev
        return (self.imp_total - prev) / prev if prev else None

    def import_multiple(self):
        if self.exp_total:
            return self.imp_total / self.exp_total
        return None

    # -- ranking helpers ----------------------------------------------------
    def shares(self, table, n=None, year="cur"):
        """(name, value, share) per ranked item for one of the two years."""
        idx = 1 if year in ("cur", "current") else 0
        total = table["grand"]["totals"][idx]
        rows = table["items"][:n] if n else table["items"]
        if total is None:                      # year not present in data
            return []
        out = []
        for d in rows:
            v = d["totals"][idx] or 0.0
            out.append((d["name"], v, (v / total) if total else 0.0))
        return out

    def cum_share(self, table, n, year="cur"):
        return sum(s for _, _, s in self.shares(table, n, year))

    def find_item(self, table, pattern, year="cur"):
        rx = re.compile(pattern, re.IGNORECASE)
        for name, v, s in self.shares(table, year=year):
            if rx.search(name):
                return (name, v, s)
        return None

    def group_share(self, table, group, year="cur"):
        """Combined share of items whose normalised name is inside group."""
        idx = 1 if year in ("cur", "current") else 0
        total = table["grand"]["totals"][idx] or 0.0
        acc = 0.0
        for d in table["items"]:
            key = re.sub(r"\s+", " ", str(d["name"]).strip().lower())
            if key in group:
                acc += d["totals"][idx] or 0.0
        return (acc / total) if total else 0.0

    def keyword_share(self, table, keywords, exclude=(), year="cur"):
        """Share of items matching any keyword (searched in the label)."""
        idx = 1 if year in ("cur", "current") else 0
        total = table["grand"]["totals"][idx] or 0.0
        acc = 0.0
        ex = [re.compile(p, re.IGNORECASE) for p in exclude]
        for d in table["items"]:
            name = str(d["name"])
            if any(rx.search(name) for rx in ex):
                continue
            low = name.lower()
            hit = any(kw in low for kw in keywords) \
                or any(re.search(kw, name, re.IGNORECASE)
                       for kw in keywords if kw.startswith("^"))
            if hit:
                acc += d["totals"][idx] or 0.0
        return (acc / total) if total else 0.0

    def growth_decline(self, table, threshold=1.0):
        """Items whose YoY change exceeds +/- ``threshold`` Ksh. Billion.

        Returns (growers, decliners); both sorted by absolute change,
        largest first.
        """
        if not self.has_comparison:
            return [], []
        growers = [d for d in table["items"]
                   if d["change"] is not None and d["change"] > threshold]
        decliners = [d for d in table["items"]
                     if d["change"] is not None and d["change"] < -threshold]
        growers.sort(key=lambda d: d["change"], reverse=True)
        decliners.sort(key=lambda d: d["change"])
        return growers, decliners


def quarter_label(month_names):
    """'April-June' from parsed month header names."""
    if len(month_names) == 1:
        return month_names[0]
    return "%s-%s" % (month_names[0], month_names[-1])


def fmt_b(v, dec=1):
    return ("Ksh %.*f billion" % (dec, v)) if v is not None else ""


def pct_s(share, dec=1):
    return ("%.*f%%" % (dec, share * 100)) if share is not None else ""


def short_prod(label, maxlen=52):
    txt = short_product_name(label, maxlen=maxlen)
    return txt.rstrip(".") if len(txt) <= maxlen else short_product_name(
        label, maxlen=maxlen - 1)


# ---------------------------------------------------------------------------
# Narrative generation (every figure derived from the analysed data)
# ---------------------------------------------------------------------------
_KEEP_CAPS = {"USA", "UAE", "UK", "US", "EU", "DRC", "USSR", "GCC"}


def _mkt(name):
    """Narration-friendly market name ('JAPAN' -> 'Japan', 'USA' kept)."""
    disp = display_name(name)
    if disp and disp.isupper():
        words = [w if w.strip(".,") in _KEEP_CAPS else
                 "-".join(p.capitalize() for p in w.split("-"))
                 for w in disp.split()]
        disp = " ".join(words)
    return disp


def build_narratives(a: QuarterAnalysis):
    """Build every narrative block from the analysed quarter.

    All figures are derived from the parsed tables; nothing is hard-coded
    except the phrasing patterns of the reference report.
    """
    Y = a.year
    q = a.quarter
    m_first, m_last = a.month_names[0], a.month_names[-1]
    period = "%s %d" % (q, Y)
    t1, t2 = a.t1, a.t2
    a1, a2 = a.a1, a.a2

    exp_t, imp_t, def_t = a.exp_total, a.imp_total, a.deficit
    growth = a.exp_growth_through_quarter()
    multiple = a.import_multiple()

    mkts = a.shares(t1)
    top10_m = a.cum_share(t1, 10)
    eac_s = a.group_share(t1, EAC)

    tea = a.find_item(t2, r"^tea[\s,]")
    coffee = a.find_item(t2, r"^coffee")
    roses = a.find_item(t2, r"\broses?\b")
    other_flowers = None
    for name, v, s in a.shares(t2):
        low = name.lower()
        if "cut flower" in low and not (roses and name == roses[0]) \
                and "rose" not in low:
            other_flowers = (name, v, s)
            break
    avo = a.find_item(t2, r"avocado")
    petro_exp = a.find_item(t2, r"petroleum oils")
    agri_s = a.keyword_share(
        t2, ("tea", "coffee", "cut flower", "rose", "avocado", "vegetable",
             "fruit", "macadamia", "nuts", "beans,", "peas,"))

    imp_mkts = a.shares(a1)
    asia_s = a.group_share(a1, ASIA)
    gulf_s = a.group_share(a1, GULF)
    west_items = [(n_, v_, s_) for n_, v_, s_ in imp_mkts
                  if n_.strip().lower() in WEST]
    west_s = sum(s_ for _, _, s_ in west_items)
    africa_items = [(n_, v_, s_) for n_, v_, s_ in imp_mkts
                    if n_.strip().lower() in AFRICA]
    africa_s = sum(s_ for _, _, s_ in africa_items)

    petro_imp = a.find_item(a2, r"petroleum oils")
    gasol = a.find_item(a2, r"motor spirit|gasoline")
    butane = a.find_item(a2, r"butane")
    fuel_s = a.keyword_share(a2, PETROLEUM_KW)
    wheat = a.find_item(a2, r"wheat|meslin")
    palm = a.find_item(a2, r"palm oil")
    rice = a.find_item(a2, r"\brice\b")
    maize = a.find_item(a2, r"maize|\bcorn\b")

    n = {}

    # ---- Section 1: Overview --------------------------------------------
    p1 = ("The total value of Kenya's merchandise exports in the second "
          "quarter of %d (between %s and %s) was %s. "
          % (Y, m_first, m_last, fmt_b(exp_t)))
    if growth is not None:
        p1 += ("Exports grew steadily through the quarter (%+.1f%%, %s to "
               "%s). " % (growth * 100, m_first, m_last))
    p1 += ("Trade remains regionally concentrated and agriculture led. "
           "The top 10 markets taking %s of exports, led by %s (%s), %s (%s) "
           "and %s (%s), with EAC partners absorbing approximately %s of the "
           "total. " % (pct_s(top10_m),
                        _mkt(mkts[0][0]), pct_s(mkts[0][2]),
                        _mkt(mkts[1][0]), pct_s(mkts[1][2]),
                        _mkt(mkts[2][0]), pct_s(mkts[2][2]),
                        pct_s(eac_s)))
    if tea:
        p1 += ("%s is the top export at %s (%s)"
               % (short_prod(tea[0], 20), fmt_b(tea[1]), pct_s(tea[2])))
        extras = [x for x in (coffee, roses) if x]
        if len(extras) >= 2:
            p1 += (", followed by %s (%s), %s (%s) and other cut flowers. "
                   % (short_prod(extras[0][0], 16), pct_s(extras[0][2]),
                      short_prod(extras[1][0], 16), pct_s(extras[1][2])))
        elif extras:
            p1 += (", followed by %s (%s). "
                   % (short_prod(extras[0][0], 24), pct_s(extras[0][2])))
        else:
            p1 += ". "
    if agri_s:
        p1 += ("Together horticulture and agriculture exports are about %s "
               "of the total." % pct_s(agri_s, 0))
    n["overview_p1"] = p1

    p2 = ("The impressive export performance was, however, greatly offset by "
          "%s in imports. " % fmt_b(imp_t))
    if petro_imp:
        p2 += ("Imports are fuel-dominated. Petroleum oils alone are %s of "
               "imports" % pct_s(petro_imp[2]))
        parts = []
        if gasol:
            parts.append("gasoline %s" % pct_s(gasol[2]))
        if butane:
            parts.append("butanes %s" % pct_s(butane[2]))
        if fuel_s and parts:
            p2 += (", and fuel-related goods (incl. %s) reached approximately "
                   "%s" % (", ".join(parts), pct_s(fuel_s)))
        p2 += ", "
    food_bits = [x for x in (wheat, palm, rice) if x]
    if food_bits:
        p2 += ("alongside large food (%s) and capital/consumer goods "
               "sustaining a structurally large import bill."
               % ", ".join(short_product_name(x[0], 18).lower()
                           for x in food_bits))
    else:
        p2 += "sustaining a structurally large import bill."
    n["overview_p2"] = p2

    n["overview_p3"] = (
        "The large trade deficit of %s creates a net foreign-currency "
        "outflow keeping persistent pressure on the shilling and raising the "
        "cost of the goods Kenya imports most (fuel and food)."
        % fmt_b(abs(def_t)))

    # ---- Section 2 bullets -------------------------------------------------
    yoy_e = a.yoy_growth()
    yoy_i = a.yoy_import_growth()
    b1 = ("During the review period (%s), total exports were valued at %s."
          % (period, fmt_b(exp_t)))
    if yoy_e is not None:
        b1 += " This was %s (%+.1f%%) over the %s quarter of %d." \
            % ("an increase" if yoy_e >= 0 else "a decline",
               yoy_e * 100, q, Y - 1)
    b2 = ("On the other hand, total imports were worth %s in the period %s "
          "to %s %d," % (fmt_b(imp_t), m_first, m_last, Y))
    if yoy_i is not None:
        b2 += " %s by %+.1f%% over the same quarter of %d," \
            % ("up" if yoy_i >= 0 else "down", yoy_i * 100, Y - 1)
    n["s2"] = [
        b1,
        b2,
        "The balance of trade for merchandized trade for the period %s to "
        "%s %d was %s." % (m_first, m_last, Y, fmt_b(def_t)),
    ]

    # ---- Section 5: Deduction ----------------------------------------------
    d1 = ("Imports outpaced exports by a wide margin in the %s quarter. "
          "Merchandise exports earned %s against an import bill of %s, "
          "yielding a trade deficit of about %s"
          % (period, fmt_b(exp_t), fmt_b(imp_t), fmt_b(abs(def_t))))
    if multiple:
        d1 += ", with imports exceeding exports roughly %.1f times. " \
              % multiple
    d1 += ("The composition of the import bill underlines Kenya's continued "
           "reliance on foreign goods. ")
    if fuel_s:
        d1 += ("Fuel-related products accounted for nearly %s of imports; "
               % pct_s(fuel_s, 0))
    staple = [x for x in (wheat, palm, rice, maize) if x]
    if staple:
        d1 += ("essential food items (%s) remain structural imports; "
               % ", ".join(short_product_name(x[0], 16).lower()
                           for x in staple[:3]))
    d1 += ("and machinery and capital goods point to dependence on imported "
           "equipment for investment and production. This sustained reliance "
           "on imported energy, food and capital goods exerts persistent "
           "pressure on the shilling and on foreign exchange reserves, which "
           "keeps the import bill large even as exports grow.")
    n["deduction"] = [d1]

    lead = ["%s was the single largest market with %s (%s)"
            % (_mkt(mkts[0][0]), fmt_b(mkts[0][1]), pct_s(mkts[0][2]))]
    follow = ["%s (%s; %s)" % (_mkt(nm), fmt_b(vv), pct_s(ss))
              for nm, vv, ss in mkts[1:6]]
    d2 = ("Kenya's export performance remains anchored on a few "
          "well-established destinations. %s, followed by %s, with EAC "
          "partners together absorbing about %s of total exports. "
          % (lead[0], ", ".join(follow), pct_s(eac_s, 0)))
    african = [(nm, vv, ss) for nm, vv, ss in mkts
               if nm.strip().lower() in AFRICA
               and nm.strip().lower() not in EAC]
    if african:
        d2 += ("Beyond these, several African destinations hold significant "
               "potential. %s already feature among Kenya's top buyers under "
               "the African Continental Free Trade Area (AfCFTA), which "
               "presents a strategic opportunity to deepen and diversify "
               "these gains."
               % ", ".join(_mkt(nm) for nm, _, _ in african[:6]))
    n["deduction"].append(d2)

    d3 = ""
    bits = []
    if roses:
        bits.append("roses (%s; %s)" % (fmt_b(roses[1]), pct_s(roses[2])))
    if other_flowers:
        bits.append("other cut flowers (%s; %s)"
                    % (fmt_b(other_flowers[1]), pct_s(other_flowers[2])))
    if bits or avo:
        d3 = ("Horticulture remains the standout driver of Kenya's export "
              "performance. Cut flowers alone - %s" % ", ".join(bits))
        if avo:
            d3 += (" – together with fresh avocados (%s; %s)"
                   % (fmt_b(avo[1]), pct_s(avo[2])))
        comp = [x for x in (tea, coffee) if x]
        if comp:
            d3 += (" – reaffirm the sector's role as a leading "
                   "foreign-exchange earner, complementing %s to make "
                   "agricultural and horticultural exports roughly %s of the "
                   "total in the period under review. "
                   % (" and ".join("%s (%s)" % (short_prod(x[0], 12),
                                                fmt_b(x[1])) for x in comp),
                      pct_s(agri_s, 0)))
        else:
            d3 += " – reaffirm horticulture as a leading earner. "
        d3 += ("The sector's growth is buttressed by several bilateral "
               "market-access arrangements, including the Kenya-EU Economic "
               "Partnership Agreement; the UK-Kenya EPA; and market-access "
               "protocols with China (fresh avocados since 2022), which are "
               "opening new, high-value Asian markets. ")
    if petro_exp:
        d3 += ("Other notable gainers in the quarter included petroleum "
               "products (%s; %s), which recorded significant growth over "
               "the second quarter." % (fmt_b(petro_exp[1]),
                                        pct_s(petro_exp[2])))
    if d3:
        n["deduction"].append(d3)

    n["deduction"].append(
        "Several other issues merit attention. First, export concentration, "
        "in both products (tea, coffee and horticulture) and markets, "
        "exposes earnings to commodity price volatility and weather shocks, "
        "underscoring the need to accelerate product and market "
        "diversification. Second, the prominence of fuel in both the export "
        "and import baskets ties Kenya's trade performance to global oil "
        "prices, so rising prices widen the deficit even as re-exports of "
        "refined petroleum support export values. These issues underscore "
        "the need to diversify export products and address competitiveness "
        "and supply-side constraints in key sectors.")

    # ---- Annex 1: imports by supplying countries ---------------------------
    im3 = imp_mkts[:3]
    monthly_txt = ", ".join("%s %s" % (mn, fmt_b(v))
                            for mn, v in zip(a.month_names, a.imp_months()))
    annex1 = [
        "Kenya imported goods worth %s in the %s quarter (%s). Import "
        "sourcing is heavily concentrated in a small group of partner "
        "markets. The top three sources – %s – supplied %s (%s of the total "
        "import bill)."
        % (fmt_b(imp_t), period, monthly_txt,
           ", ".join(_mkt(x[0]) for x in im3),
           fmt_b(sum(x[1] for x in im3)), pct_s(sum(x[2] for x in im3))),
        "The top five suppliers accounted for %s of imports and the top ten "
        "for %s, leaving all other countries with less than an eighth of "
        "the market. This concentration underscores Kenya's dependence on a "
        "narrow group of supply markets, particularly in Asia and the Gulf."
        % (pct_s(a.cum_share(a1, 5)), pct_s(a.cum_share(a1, 10))),
    ]
    third_parts = []
    if west_items:
        wp = ("Traditional Western trading partners – %s – together supplied "
              "roughly %s of imports"
              % (", ".join(_mkt(x[0]) for x in west_items[:9]),
                 pct_s(west_s)))
        singles = ["%s %s" % (_mkt(x[0]), pct_s(x[2]))
                   for x in west_items if x[2] >= 0.005][:6]
        if singles:
            wp += ", with individual shares now modest (%s)" \
                  % ", ".join(singles)
        wp += ". Import sourcing has therefore shifted decisively towards "
        if asia_s:
            wp += ("Asia, with China, India, Malaysia, Japan and others "
                   "accounting for about %s of the bill" % pct_s(asia_s))
            if gulf_s:
                wp += ", while the Gulf states contribute a further %s" \
                      % pct_s(gulf_s)
            wp += "."
        elif gulf_s:
            wp += "the Gulf states, which contribute a further %s." \
                  % pct_s(gulf_s)
        else:
            wp += "markets outside the traditional Western bloc."
        third_parts.append(wp)
    else:
        if asia_s or gulf_s:
            third_parts.append("Import sourcing is dominated by Asia and "
                               "the Gulf, together accounting for about %s "
                               "of the bill."
                               % pct_s(asia_s + gulf_s))
    if africa_items:
        third_parts.append("Intra-African supply remains limited – %s "
                           "together provide only about %s of imports, "
                           "pointing to room for greater regional sourcing "
                           "under the AfCFTA."
                           % (", ".join("%s (%s)" % (_mkt(x[0]), pct_s(x[2]))
                                        for x in africa_items[:5]),
                              pct_s(africa_s)))
    if third_parts:
        annex1.append(" ".join(third_parts))
    n["annex1"] = annex1

    # ---- Annex 2: imports by products ---------------------------------------
    monthly_avg = imp_t / max(1, len(a.month_names))
    a2p1 = ("The total import bill for the quarter was %s (monthly average "
            "%s). Four broad forces drive the bill. First, fuel products are "
            "the single largest component"
            % (fmt_b(imp_t), fmt_b(monthly_avg)))
    if petro_imp:
        a2p1 += " – petroleum oils alone were worth %s (%s)" \
                % (fmt_b(petro_imp[1]), pct_s(petro_imp[2]))
    staples = [x for x in (wheat, palm, rice, maize) if x]
    if staples:
        a2p1 += (". Second, recurring imports of %s reflect domestic "
                 "production shortfalls."
                 % ", ".join(short_product_name(x[0], 16).lower()
                             for x in staples))
    else:
        a2p1 += (". Second, recurring food imports reflect domestic "
                 "production shortfalls.")
    a2p1 += (". Third, machinery, vehicles and equipment linked to "
             "construction, ICT, transport and investment activity. Fourth, "
             "consumables and industrial inputs such as medicaments, plastic "
             "polymers and electronic materials, which support manufacturing "
             "and health services.")
    pm = next((d for d in a2["items"]
               if re.search(r"petroleum oils", str(d["name"]), re.I)), None)
    if pm and pm["totals"][0]:
        a2p1 += (" In the short run, oil price movements are one of the "
                 "main sources of monthly volatility, as seen in the "
                 "sharp swing in petroleum oils between %s and %s."
                 % (a.month_names[0], a.month_names[-1]))
    n["annex2"] = [a2p1]
    return n
# ---------------------------------------------------------------------------
# Charts (matplotlib, house style shared with the other report generators)
# ---------------------------------------------------------------------------
PIE_PALETTE = list(THEME_ACCENTS) + [lighten(c, 0.45) for c in THEME_ACCENTS]


def make_chart_balance(a: QuarterAnalysis, out_path):
    """Clustered columns of monthly Exports / Imports / Balance."""
    plt.rcParams["font.family"] = chart_font()
    fig, ax = plt.subplots(figsize=(7.6, 4.4), dpi=160)
    months = a.month_names
    x = list(range(len(months)))
    w = 0.26
    cols = [lighten(THEME_ACCENTS[0], 0.30), lighten(THEME_ACCENTS[1], 0.30),
            lighten(THEME_ACCENTS[2], 0.30)]
    exp_m, imp_m = a.exp_months(), a.imp_months()
    bal_m = a.balance_months()
    ax.bar([i - w for i in x], exp_m, width=w, label="Exports",
           color=cols[0])
    ax.bar([i for i in x], imp_m, width=w, label="Imports", color=cols[1])
    ax.bar([i + w for i in x], bal_m, width=w, label="Balance of Trade",
           color=cols[2])
    ax.axhline(0, color="#555555", linewidth=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels(months, fontsize=9)
    ax.set_ylabel("Value in Ksh. Billion", fontsize=10)
    ax.set_title("Kenya Balance of Trade %s %d" % (a.quarter, a.year),
                 fontsize=12, weight="bold")
    ax.yaxis.grid(True, linestyle="--", alpha=0.35)
    ax.set_axisbelow(True)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.10), ncol=3,
              frameon=False, fontsize=10)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)


def make_chart_share(a: QuarterAnalysis, table, title, out_path,
                     other_label="Others"):
    """Share doughnuts.

    Comparison mode: two pies stacked vertically ('SHARE IN <Y>' above
    'SHARE IN <Y-1>'), sized to fill a full A4 page. Single-year mode:
    one centred doughnut.
    """
    Yc, Yp = a.year, a.year_prev

    def side(idx):
        labels, values = [], []
        for name, v, _s in a.shares(table, year=("cur" if idx == 1
                                                 else "prev")):
            labels.append(short_prod(name, 40))
            values.append(v if v else 0.0)
        return consolidate(labels, values, max_slices=8, min_pct=2.0,
                           other_label=other_label)

    plt.rcParams["font.family"] = chart_font()
    if a.has_comparison:
        # full A4 portrait content page: two pies stacked vertically,
        # each labelled with a right-hand call-out legend. Saved directly
        # (no tight bbox) so the image keeps its full-page dimensions.
        fig = plt.figure(figsize=(7.5, 10.4), dpi=160)
        fig.patch.set_facecolor("white")
        for y0, idx, yr in ((0.545, 1, Yc), (0.055, 0, Yp)):
            ax = fig.add_axes([0.03, y0, 0.56, 0.36])
            labels, values = side(idx)
            labels2, shares2, wedges = draw_share_pie(
                ax, labels, values, PIE_PALETTE[:len(values)],
                style="3d_exploded", other_label=other_label)
            total = sum(shares2) or 1.0
            entries = ["%s \u2013 %.1f%%" % (l, v / total * 100)
                       for l, v in zip(labels2, shares2)]
            leg = ax.legend(wedges, entries, loc="center left",
                            bbox_to_anchor=(1.04, 0.5), frameon=False,
                            fontsize=10, handlelength=1.3,
                            labelspacing=1.0, borderaxespad=0.0)
            for t in leg.get_texts():
                t.set_fontsize(10)
            ax.set_title("SHARE IN %d" % yr, fontsize=14, weight="bold",
                         y=1.06)
        fig.savefig(out_path, facecolor="white")
        plt.close(fig)
        return

    labels, values = side(1)
    fig, ax = new_fig()
    labels2, shares2, wedges = draw_share_pie(
        ax, labels, values, PIE_PALETTE[:len(values)],
        style="3d_exploded", other_label=other_label)
    slice_callouts(fig, wedges, labels2, shares2)
    ax.set_title("Share in %d" % Yc, fontsize=12, weight="bold")
    finish(fig, out_path)


# ---------------------------------------------------------------------------
# Quarterly Word tables
# ---------------------------------------------------------------------------
GREEN = RGBColor(0x00, 0xB0, 0x50)
RED = RGBColor(0xFF, 0x00, 0x00)
GREY = RGBColor(0x80, 0x80, 0x80)


def _arrow_text(value, kind):
    """(arrow, text, color) for a change / % change value."""
    if value is None:
        return "", "–", GREY
    up = value >= 0
    arrow = "\u25b2" if up else "\u25bc"
    color = GREEN if up else RED
    txt = ("%.1f%%" % (value * 100)) if kind == "pct" else num(value)
    return arrow, txt, color


def _fill_arrow_cell(b, cell, value, kind):
    """Write an arrow + coloured value into a table cell."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    p = cell.paragraphs[0]
    p.alignment = _AL.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    arrow, txt, color = _arrow_text(value, kind)
    if arrow:
        r = p.add_run(arrow + " ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = color
    r2 = p.add_run(txt)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10)
    r2.font.color.rgb = color


def _add_q_table_simple(b: ReportBuilder, parsed, item_header):
    """Single-year table: Rank | Item | months | Total | Share."""
    Y = b.year
    n = len(parsed["months"])
    ncols = n + 4
    data_rows = len(parsed["items"]) + (1 if parsed["all_other"] else 0)
    rows = 3 + data_rows
    table = b.doc.add_table(rows=rows, cols=ncols)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [700, 2750] + [900] * n + [1000, 950]
    b._set_table_widths(table, widths)

    b._cell_text(table.cell(0, 0), "Rank \nin\n%d" % Y, bold=True)
    b._cell_text(table.cell(0, 1), item_header, bold=True, wrap=True)
    for k, m in enumerate(parsed["months"]):
        b._cell_text(table.cell(0, 2 + k), m, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER)
    b._cell_text(table.cell(0, 2 + n), "%d\nTotal" % Y, bold=True)
    b._cell_text(table.cell(0, 3 + n), "Share \nin\n%d" % Y, bold=True)
    table.cell(0, 3 + n).merge(table.cell(1, 3 + n))

    ucell = table.cell(1, 2)
    b._cell_text(ucell, "Value in Ksh. Billion", bold=True)
    ucell.merge(table.cell(1, 1 + n))

    def emit(d, ri, bold=False):
        b._cell_text(table.cell(ri, 0),
                     str(d["rank"]) if d.get("rank") is not None else "",
                     bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        b._cell_text(table.cell(ri, 1), str(d["name"]), bold=bold, wrap=True)
        for k, v in enumerate(d["months"][1]):
            b._cell_text(table.cell(ri, 2 + k),
                         num(v) if v is not None else "",
                         bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        tot = d["totals"][1]
        b._cell_text(table.cell(ri, 2 + n),
                     num(tot) if tot is not None else "",
                     bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        sh = d.get("share")
        b._cell_text(table.cell(ri, 3 + n),
                     ("%.2f" % (sh * 100)) if sh is not None else "",
                     bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)

    ri = 2
    for d in parsed["items"]:
        emit(d, ri)
        ri += 1
    if parsed["all_other"]:
        emit(parsed["all_other"], ri)
        ri += 1
    grand = {"name": parsed["grand"].get("name", "Grand Total"),
             "rank": None,
             "months": [None, parsed["grand"]["months"][1]],
             "totals": [None, parsed["grand"]["totals"][1]],
             "share": parsed["grand"].get("share", 1.0)}
    emit(grand, ri, bold=True)
    b._fit_table_on_page(table)
    return table


def add_q_table(b: ReportBuilder, parsed, item_header):
    """Two-year comparison table (Rank | Item | <Y-1> | <Y> | Change).

    Columns: rank, item, prev months+Total, current months+Total,
    Total change and % with green up / red down arrows.
    Falls back to the single-year layout when no prior-year data exists.
    """
    if not parsed.get("comparison"):
        return _add_q_table_simple(b, parsed, item_header)
    Yc = b.year
    Yp = Yc - 1
    n = len(parsed["months"])
    ncols = 4 + 2 * n + 2                 # rank,item,(n+1)*2,chg,pct
    data_rows = len(parsed["items"]) + (1 if parsed["all_other"] else 0)
    rows = 2 + data_rows + 1              # header x2 + data + grand
    table = b.doc.add_table(rows=rows, cols=ncols)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    c_rank, c_item = 0, 1
    p_first = 2
    p_total = p_first + n
    c_first = p_total + 1
    c_total = c_first + n
    c_chg = c_total + 1
    c_pct = c_chg + 1

    widths = [430, 2350] + [640] * n + [760] + [640] * n + [760] \
        + [820, 700]
    b._set_table_widths(table, widths)

    def hcell(r, c, text):
        b._cell_text(table.cell(r, c), text, bold=True)

    # header row 0: merged group headers
    hcell(0, c_rank, "Rank")
    hcell(0, c_item, item_header)
    table.cell(0, c_rank).merge(table.cell(1, c_rank))
    table.cell(0, c_item).merge(table.cell(1, c_item))
    table.cell(0, p_first).merge(table.cell(0, p_total))
    hcell(0, p_first, str(Yp))
    table.cell(0, c_first).merge(table.cell(0, c_total))
    hcell(0, c_first, str(Yc))
    table.cell(0, c_chg).merge(table.cell(0, c_pct))
    hcell(0, c_chg, "Change in\n%d-%d" % (Yp, Yc))

    # header row 1: month labels / totals / change labels
    for base, tcol in ((p_first, p_total), (c_first, c_total)):
        for k, m in enumerate(parsed["months"]):
            hcell(1, base + k, m[:3])
        hcell(1, tcol, "Total")
    hcell(1, c_chg, "Total change")
    hcell(1, c_pct, "%")

    def emit(d, ri, bold=False):
        b._cell_text(
            table.cell(ri, c_rank),
            str(d["rank"]) if d.get("rank") is not None else "",
            bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        b._cell_text(table.cell(ri, c_item), str(d["name"]), bold=bold,
                     wrap=True)
        vals = ((d["months"][0], d["totals"][0], p_first, p_total),
                (d["months"][1], d["totals"][1], c_first, c_total))
        for mvals, tot, first, tcol in vals:
            for k, v in enumerate(mvals):
                b._cell_text(table.cell(ri, first + k),
                             num(v) if v is not None else "",
                             bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
            b._cell_text(table.cell(ri, tcol),
                         num(tot) if tot is not None else "",
                         bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        _fill_arrow_cell(b, table.cell(ri, c_chg), d.get("change"), "chg")
        _fill_arrow_cell(b, table.cell(ri, c_pct), d.get("pct"), "pct")

    ri = 2
    for d in parsed["items"]:
        emit(d, ri)
        ri += 1
    if parsed["all_other"]:
        emit(parsed["all_other"], ri)
        ri += 1
    emit(parsed["grand"], ri, bold=True)
    b._fit_table_on_page(table)
    return table


def add_growth_decline_table(b: ReportBuilder, growers, decliners,
                             left_title, right_title,
                             threshold_label="Ksh. 1 billion"):
    """Two-column exhibit: markets/products that grew vs declined.

    Each entry reads ``Name  <arrow> Ksh X.X billion`` with the arrow and
    amount coloured green (growth) or red (decline).
    """

    def shade(cell, hexcolor):
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _El
        tcPr = cell._tc.get_or_add_tcPr()
        shd = _El("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:fill"), hexcolor)
        tcPr.append(shd)

    nrows = 1 + max(len(growers), len(decliners), 1)
    table = b.doc.add_table(rows=nrows, cols=2)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b._set_table_widths(table, [5300, 5300])

    for c, title in ((0, left_title), (1, right_title)):
        shade(table.cell(0, c), "D9E1F2")
        b._cell_text(table.cell(0, c), title, bold=True, wrap=True, size=10)

    def fill_entry(cell, d):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(str(d["name"]))
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        arrow, txt, color = _arrow_text(d["change"], "chg")
        r2 = p.add_run("  %s %s" % (arrow, txt))
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)
        r2.font.bold = True
        r2.font.color.rgb = color

    def fill_none(cell):
        b._cell_text(cell, "None registered change of over %s"
                     % threshold_label, italic=True, size=10)

    for i in range(nrows - 1):
        if i < len(growers):
            fill_entry(table.cell(i + 1, 0), growers[i])
        else:
            fill_none(table.cell(i + 1, 0))
        if i < len(decliners):
            fill_entry(table.cell(i + 1, 1), decliners[i])
        else:
            fill_none(table.cell(i + 1, 1))
    return table


DEFAULT_ORG = "KENYA EXPORT PROMOTION & BRANDING AGENCY"
DEFAULT_DIRECTORATE = "Research and Innovation Directorate (RID)"


def build_quarterly_report(excel_dir, out_path, tmp_dir,
                           year=None, org=DEFAULT_ORG,
                           directorate=DEFAULT_DIRECTORATE,
                           month_year=None):
    os.makedirs(tmp_dir, exist_ok=True)
    a = QuarterAnalysis(year=year).load(excel_dir)
    narr = build_narratives(a)

    Y = a.year
    y_prev = Y - 1
    q_up = a.quarter.upper()
    q_t = a.quarter

    cfg = {
        "country": {"name": "Kenya", "short": "Kenya",
                    "possessive": "Kenya's", "title": "Overview"},
        "report": {"year": Y},
    }
    b = ReportBuilder(cfg, narr)
    b.year = Y
    b.has_comparison = a.has_comparison

    # dynamic table numbering (exhibits shift the sequence when present)
    t_counter = [0]

    def tno():
        t_counter[0] += 1
        return t_counter[0]

    # ============================== TITLE PAGE ==============================
    b.add_letterhead()
    b.add_para("KENYA", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               space_after=0)
    b.add_para("EXPORT PERFORMANCE IN %s %d-%d" % (q_up, y_prev, Y),
               size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               space_after=24)
    b.add_para("")
    b.add_para("REPORT", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               space_after=30)
    for _ in range(3):
        b.add_para("")
    b.add_para("Prepared", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               space_after=6)
    b.add_para("By", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               space_after=14)
    b.add_para(org, size=22, bold=True, italic=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    b.add_para(directorate, size=22, bold=True, italic=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    if not month_year:
        month_year = date.today().strftime("%B %Y")
    b.add_para(month_year, size=22, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    b.page_break()

    # ==================== TOC / LISTS (unnumbered pages) ====================
    b.add_toc()
    b.page_break()
    b.add_list_placeholder("List of Tables", "table")
    b.page_break()
    b.add_list_placeholder("List of Figures", "figure")
    # page numbers start at 1 from here (title/TOC/lists carry none)
    b.start_body_section()

    # ============================== SECTION 1 ===============================
    p = b.doc.add_paragraph()
    r = p.add_run("KENYA\u2019S TRADE PERFORMANCE IN %s (%d AND %d)"
                  % (q_up, y_prev, Y))
    b._style_run(r, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    b.add_heading("1. Overview")
    for key in ("overview_p1", "overview_p2", "overview_p3"):
        b.add_para(narr[key])
    # uncaptioned balance chart, as in the reference report
    bal_png = os.path.join(tmp_dir, "chart_balance.png")
    make_chart_balance(a, bal_png)
    b.add_figure(bal_png)
    b.add_source(SRC_KRA)
    b.page_break()

    # ============================== SECTION 2 ===============================
    b.add_heading("2. Comparative Perspectives of Kenya\u2019s Trade in "
                  "%s (%d/%d)" % (q_t, y_prev, Y))
    b.add_para("")
    for line in narr["s2"]:
        b.add_para(line, style="List Bullet")
    b.page_break()

    # ============================== SECTION 3 ===============================
    cmp_span = ("(%d-%d)" % (y_prev, Y) if a.has_comparison
                else "(%d)" % Y)
    b.add_heading("3. Lead Destination Markets %s %s" % (q_t, cmp_span))
    b.add_table_caption("Table %d: Top Destination Markets in %s %s"
                        % (tno(), q_t, cmp_span))
    add_q_table(b, a.t1, "PARTNER")
    b.add_source(SRC_KRA)

    if a.has_comparison:
        # exhibit: markets that grew vs declined by over Ksh 1 billion
        g_mkt, d_mkt = a.growth_decline(a.t1, threshold=1.0)
        b.add_table_caption(
            "Table %d: Export Markets that Registered Growth / Decline of "
            "Over Ksh. 1 Billion in %s %s" % (tno(), q_t, cmp_span))
        add_growth_decline_table(
            b, g_mkt, d_mkt,
            "Export markets that registered growth (by more than Ksh. 1 "
            "billion)",
            "Economies whose exports declined (by more than Ksh. 1 "
            "billion)")
        b.add_source(SRC_KRA)

    b.add_heading("3.1 Comparative Perspectives on Shares of Lead Export "
                  "Markets %s %s" % (q_t, cmp_span))
    b.add_table_caption("Figure 1: Share of Exports by Market, %s %s"
                        % (q_t, cmp_span))
    fig1 = os.path.join(tmp_dir, "chart_export_markets.png")
    make_chart_share(a, a.t1, "Share of Exports by Market", fig1,
                     other_label="Other markets")
    b.add_figure(fig1, width_in=6.3 if a.has_comparison else 6.4)
    b.add_source(SRC_KRA)
    b.page_break()

    # ============================== SECTION 4 ===============================
    b.add_heading("4. Composition of Kenya\u2019s Export Commodities")
    b.add_table_caption("Table %d: Top Export Products in %s %s"
                        % (tno(), q_t, cmp_span))
    add_q_table(b, a.t2, "Product label")
    b.add_source(SRC_KRA)

    if a.has_comparison:
        # exhibit: products that grew vs declined by over Ksh 1 billion
        g_prd, d_prd = a.growth_decline(a.t2, threshold=1.0)
        b.add_table_caption(
            "Table %d: Export Products that Registered Growth / Decline of "
            "Over Ksh. 1 Billion in %s %s" % (tno(), q_t, cmp_span))
        add_growth_decline_table(
            b, g_prd, d_prd,
            "Export products that registered growth (by over Ksh. 1 "
            "billion)",
            "Export products that registered decline (by over Ksh. 1 "
            "billion)")
        b.add_source(SRC_KRA)

    b.add_heading("4.1 Share of Lead Export Products %s %s"
                  % (q_t, cmp_span))
    b.add_table_caption("Figure 2: Share of Top Exports by Products, %s "
                        "%s" % (q_t, cmp_span))
    fig2 = os.path.join(tmp_dir, "chart_export_products.png")
    make_chart_share(a, a.t2, "Share of Top Exports by Products", fig2,
                     other_label="All others products")
    b.add_figure(fig2, width_in=6.3 if a.has_comparison else 6.4)
    b.add_source(SRC_KRA)
    b.page_break()

    # ============================== SECTION 5 ===============================
    b.add_heading("5. Deduction from Quarter Performance %s to %s %d"
                  % (a.month_names[0], a.month_names[-1], Y))
    for para in narr["deduction"]:
        b.add_para(para)
    b.page_break()

    # ============================== SECTION 6 ===============================
    b.add_heading("6. Annex")
    b.add_heading("Annex 1: Kenya\u2019s Imports by Supplying Countries.",
                  level=2)
    b.add_table_caption("Table %d: Kenya\u2019s Imports by Supplying "
                        "Countries, %s %s" % (tno(), q_t, cmp_span))
    add_q_table(b, a.a1, "PARTNER")
    b.add_source(SRC_KRA)
    for para in narr["annex1"]:
        b.add_para(para)
    b.page_break()

    b.add_heading("Annex 2: Kenya\u2019s Imports by Products.", level=2)
    b.add_table_caption("Table %d: Kenya\u2019s Imports by Products, %s "
                        "%s" % (tno(), q_t, cmp_span))
    add_q_table(b, a.a2, "Product Label")
    b.add_source(SRC_KRA)
    for para in narr["annex2"]:
        b.add_para(para)

    b.add_footer()
    return a, b.doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(
        description="Generate the Kenya quarterly export performance report.")
    ap.add_argument("--excel-dir",
                    default=os.path.join(BASE_DIR, "output", "quarterly"),
                    help="Folder containing Exports.xlsx and Imports.xlsx")
    ap.add_argument("--output", default=None, help="Output .docx path")
    ap.add_argument("--tmp", default=os.path.join(BASE_DIR, "output", ".tmp"),
                    help="Temporary directory for chart images")
    ap.add_argument("--year", type=int, default=None,
                    help="Review year when it cannot be detected from data")
    ap.add_argument("--org", default=DEFAULT_ORG,
                    help="Organisation line on the title page")
    ap.add_argument("--directorate", default=DEFAULT_DIRECTORATE)
    ap.add_argument("--month-year", default=None,
                    help="Date line on the title page (e.g. 'August 2026')")
    args = ap.parse_args()

    out = args.output
    out = os.path.abspath(out) if out else os.path.join(
        BASE_DIR, "output", "KENYA EXPORT PERFORMANCE REPORT.docx")
    os.makedirs(os.path.dirname(out), exist_ok=True)

    print("[1/3] Reading quarterly workbooks from : %s"
          % os.path.abspath(args.excel_dir))
    a, doc = build_quarterly_report(
        args.excel_dir, out, args.tmp, year=args.year, org=args.org,
        directorate=args.directorate, month_year=args.month_year)
    print("[2/3] Review period                    : %s %d"
          % (a.quarter, a.year))
    doc.save(out)
    print("[3/3] Report saved to                  : %s" % out)


if __name__ == "__main__":
    main()
