#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_services_report.py
===========================

Generate a professional KEPROBA-style "Kenya -- <Country> Services Trade
Flow" Word report (.docx) from:

  1. Five Excel files produced by ``make_services_tables.py``:
        - "Table 1 Top Global Service Exporters.xlsx"
        - "Table 2 Top Global Service Importers.xlsx"
        - "Table 3 Kenya Service Exports.xlsx"
        - "Table 4 Kenya Service Imports.xlsx"
        - "Figure 1 Services Balance.xlsx"

  2. A per-country JSON configuration file (same format as the goods
     pipeline).

Usage
-----
    python3 generate_services_report.py \
        --excel-dir output/service_tables \
        --config config/saudi_arabia.json \
        --output output/KENYA-SAUDI ARABIA SERVICES TRADE FLOW.docx

Dependencies: python-docx, openpyxl, matplotlib
"""

import argparse
import json
import os
import sys

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import openpyxl

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Import shared infrastructure from the goods report generator
# ---------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE_DIR)

from generate_report import (
    ReportBuilder, load_config,
    to_float, num, pct, clean_label, short_label, lighten,
    chart_font, THEME_ACCENTS, SRC, FONT_PREFERENCE,
)

# Module-level cfg_path (used by ReportBuilder.resolve_path for images)
cfg_path = ""


# ---------------------------------------------------------------------------
# Excel reading
# ---------------------------------------------------------------------------
def find_service_excel_files(excel_dir):
    """Locate the five service Excel files by keyword in their filenames."""
    found = {}
    for fname in sorted(os.listdir(excel_dir)):
        low = fname.lower()
        if not low.endswith((".xlsx", ".xlsm")):
            continue
        path = os.path.join(excel_dir, fname)
        if "table 1" in low and "service" in low:
            found["table1"] = path
        elif "table 2" in low and "service" in low:
            found["table2"] = path
        elif "table 3" in low and "kenya" in low:
            found["table3"] = path
        elif "table 4" in low and "kenya" in low:
            found["table4"] = path
        elif "balance" in low or "figure 1" in low:
            found["balance"] = path
    # Fallback: try broader matching
    if "table1" not in found:
        for fname in sorted(os.listdir(excel_dir)):
            low = fname.lower()
            if not low.endswith((".xlsx", ".xlsm")):
                continue
            path = os.path.join(excel_dir, fname)
            if "table 1" in low:
                found.setdefault("table1", path)
            elif "table 2" in low:
                found.setdefault("table2", path)
            elif "table 3" in low:
                found.setdefault("table3", path)
            elif "table 4" in low:
                found.setdefault("table4", path)
            elif "balance" in low or "figure 1" in low:
                found.setdefault("balance", path)
    missing = [k for k in ("table1", "table2", "table3", "table4", "balance")
               if k not in found]
    if missing:
        sys.exit(f"[ERROR] Missing Excel file(s) in '{excel_dir}': {', '.join(missing)}")
    return found


def _cell_grid(ws):
    return [[c.value for c in row] for row in ws.iter_rows()]


def _find_year_run(grid):
    """Return (row_index, start_col, [years]) for the longest run of
    consecutive year integers anywhere in the grid."""
    best = None
    for ri, row in enumerate(grid):
        cur = []
        prev = None
        for ci, v in enumerate(row):
            iv = int(v) if isinstance(v, (int, float)) and not isinstance(v, bool) else None
            if iv is not None and 1990 <= iv <= 2100 and abs(v - iv) < 1e-9:
                if prev is not None and iv == prev + 1:
                    cur.append((ci, iv))
                else:
                    if 4 <= len(cur) <= 10 and (best is None or len(cur) > len(best[2])):
                        best = (ri, cur[0][0], [y for _, y in cur])
                    cur = [(ci, iv)]
                prev = iv
            else:
                prev = None
        if 4 <= len(cur) <= 10 and (best is None or len(cur) > len(best[2])):
            best = (ri, cur[0][0], [y for _, y in cur])
    return best if best else (None, None, [])


def parse_rank_table(path):
    """Parse a generic rank table (Tables 1-4)."""
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    grid = _cell_grid(ws)

    year_row_idx, year_start_col, year_header = _find_year_run(grid)
    if year_row_idx is None:
        raise ValueError(f"Could not find a consecutive year header row in {path}")
    n = len(year_header)
    val_cols = list(range(year_start_col, year_start_col + n))
    share_col = year_start_col + n

    has_code = False
    for ri in range(year_row_idx + 1):
        for v in grid[ri]:
            if isinstance(v, str) and "code" in v.strip().strip("'\"`").lower():
                has_code = True
                break
        if has_code:
            break

    name_col = 2 if has_code else 1

    data = []
    for ri in range(year_row_idx + 1, len(grid)):
        row = grid[ri]
        if not any(v is not None and str(v).strip() != "" for v in row):
            continue
        rank = int(row[0]) if isinstance(row[0], (int, float)) else None
        name = row[name_col] if name_col < len(row) else None
        code = row[1] if has_code and len(row) > 1 else None
        label = row[2] if has_code and len(row) > 2 else None
        if name is None and label is not None:
            name = label
        if name is None or (isinstance(name, str) and not name.strip()):
            continue
        name = str(name).strip()
        years = [to_float(row[c]) if c < len(row) else None for c in val_cols]
        share = to_float(row[share_col]) if share_col < len(row) else None

        if name.lower().startswith("all other"):
            kind = "all_other"
        elif name.lower() in ("world", "all services"):
            kind = "total"
        else:
            kind = "item"
        data.append({"rank": rank, "name": name, "code": code, "label": label,
                     "years": years, "share": share, "kind": kind})

    items = [d for d in data if d["kind"] == "item"]
    total = next((d for d in data if d["kind"] == "total"), None)
    all_other = next((d for d in data if d["kind"] == "all_other"), None)
    return {"has_code": has_code, "data": data, "items": items,
            "total": total, "all_other": all_other,
            "year_start_col": year_start_col, "years": year_header}


def parse_balance(path):
    """Parse Figure 1 Services Balance workbook."""
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    grid = _cell_grid(ws)

    for ri, row in enumerate(grid):
        if not row or not isinstance(row[0], str):
            continue
        lbl = row[0].strip()
        if lbl == "Exports" and ri + 2 < len(grid):
            l1 = grid[ri + 1][0].strip() if isinstance(grid[ri + 1][0], str) else ""
            l2 = grid[ri + 2][0].strip() if isinstance(grid[ri + 2][0], str) else ""
            if l1 == "Imports" and l2 == "Balance of Trade":
                years = [int(v) for v in grid[ri - 1][1:]
                         if isinstance(v, (int, float)) and int(v) >= 2000]
                if not years:
                    continue
                exports = [to_float(row[1 + i]) for i in range(len(years))]
                imports = [to_float(grid[ri + 1][1 + i]) for i in range(len(years))]
                balance = [to_float(grid[ri + 2][1 + i]) for i in range(len(years))]
                return {"years": years, "exports": exports,
                        "imports": imports, "balance": balance}
    raise ValueError(f"Could not parse balance series from {path}")


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------
class ServicesAnalysis:
    """Holds every statistic needed by the narrative and the charts."""

    def __init__(self, cfg):
        self.cfg = cfg
        self.country = cfg["country"]["name"]
        self.short = cfg["country"]["short"]
        self.possessive = cfg["country"]["possessive"]
        self.year = int(cfg["report"].get("year", 0))  # Will be overridden by data
        self.years = []
        self.iy = 0

    def load(self, excel_dir):
        files = find_service_excel_files(excel_dir)
        self.table1 = parse_rank_table(files["table1"])
        self.table2 = parse_rank_table(files["table2"])
        self.table3 = parse_rank_table(files["table3"])
        self.table4 = parse_rank_table(files["table4"])
        self.balance = parse_balance(files["balance"])
        if self.table1["years"]:
            self.years = self.table1["years"]
        # Always use the latest year from the data, regardless of config
        self.year = self.years[-1] if self.years else 2024
        self.iy = self.years.index(self.year) if self.year in self.years else len(self.years) - 1
        return self

    def top(self, table, n):
        return table["items"][:n]

    def country_top_markets_title_n(self, table):
        n = 0
        for d in table["items"]:
            if d.get("rank") is not None and d["name"].lower() != "kenya":
                n += 1
        return n


# ---------------------------------------------------------------------------
# Narrative text generation
# ---------------------------------------------------------------------------
def build_narratives(a: ServicesAnalysis, cfg):
    c = cfg["country"]
    Y = a.year
    t = {}

    def top_phrase(items, fmt, n=None):
        rows = items if n is None else items[:n]
        if not rows:
            return ""
        names = [fmt(d) for d in rows]
        if len(names) == 1:
            return names[0]
        if len(names) == 2:
            return f"{names[0]} and {names[1]}"
        return ", ".join(names[:-1]) + f", and {names[-1]}"

    def research(what):
        return f"[RESEARCH NEEDED: {what}]"

    # ---- Section 1 background ----
    t["background"] = [
        [(f"{c['name']} has its capital at {c['capital']}. ", False),
         (research("country overview - geography, total area, borders, population "
                   "size and growth, demographics, urbanisation, life expectancy"), True)],
        [(research("economy - GDP and GDP per capita (current USD), real GDP growth, "
                   "key economic sectors, income classification (e.g. high-income)"), True)],
        [(research("outlook - IMF World Economic Outlook and World Bank growth/inflation "
                   "projections, fiscal and public-debt position, policy priorities"), True)],
        [(research("trade and development policy - WTO accession date, main trade agreements, "
                   "national development/export-diversification strategy"), True)],
    ]

    # ---- Section 2: Global Services Trade ----
    exp1 = a.top(a.table1, 5)
    t["s2"] = [
        f"Global services trade reached significant levels by {Y}, with the top "
        f"service exporters continuing to dominate the market.",
    ]

    def _fmt_billion(d):
        v = d["years"][a.iy] if d["years"] and a.iy < len(d["years"]) else None
        if v is not None:
            return f"{d['name']} (USD {v:,.1f} billion)"
        return d["name"]

    def _fmt_million(d):
        v = d["years"][a.iy] if d["years"] and a.iy < len(d["years"]) else None
        if v is not None:
            return f"{d['label']} (USD {v:,.1f} million)"
        return d.get("label", d.get("name", ""))

    t["s2_exporters"] = [
        f"The leading service exporters in {Y} were: "
        f"{top_phrase(exp1, _fmt_billion)}. " if exp1 else "",
    ]

    # ---- Section 3: Top Importers ----
    imp1 = a.top(a.table2, 5)
    t["s3_importers"] = [
        f"The leading service importers in {Y} were: "
        f"{top_phrase(imp1, _fmt_billion)}. " if imp1 else "",
    ]

    # ---- Section 3.1: Kenya's service exports ----
    ke_items = a.top(a.table3, 5)
    t["s3_kenya_exports"] = [
        f"Kenya's service exports by category show the composition of the "
        f"country's services trade with the world.",
        f"The top service export categories in {Y} were: "
        f"{top_phrase(ke_items, _fmt_million)}. " if ke_items else "",
    ]

    # ---- Section 5: Balance of Trade ----
    bal = a.balance
    if bal["exports"] and bal["imports"]:
        latest_exp = bal["exports"][-1]
        latest_imp = bal["imports"][-1]
        latest_bal = bal["balance"][-1] if bal["balance"] else None
        first_exp = bal["exports"][0]
        first_imp = bal["imports"][0]
        first_year = bal["years"][0]
        t["s5"] = [
            f"Kenya's services exports grew from USD {first_exp:,.1f} million in "
            f"{first_year} to USD {latest_exp:,.1f} million in {bal['years'][-1]}.",
            f"Services imports grew from USD {first_imp:,.1f} million in "
            f"{first_year} to USD {latest_imp:,.1f} million in {bal['years'][-1]}.",
            f"Kenya maintained a positive services trade balance in {bal['years'][-1]} "
            f"of USD {latest_bal:,.1f} million."
            if latest_bal and latest_bal > 0 else
            f"Kenya recorded a services trade deficit of USD {abs(latest_bal):,.1f} "
            f"million in {bal['years'][-1]}."
            if latest_bal and latest_bal < 0 else "",
        ]
    else:
        t["s5"] = []

    # Kenya imports by category
    ki_items = a.top(a.table4, 5)
    t["s4_imports"] = [
        f"The top service import categories for Kenya in {Y} were: "
        f"{top_phrase(ki_items, _fmt_million)}. " if ki_items else "",
    ]

    return t


# ---------------------------------------------------------------------------
# Chart generation (matplotlib)
# ---------------------------------------------------------------------------
def make_chart_balance(a: ServicesAnalysis, out_path):
    """Clustered column chart of services Exports / Imports / Balance."""
    bal = a.balance
    years = bal["years"]
    exports = bal["exports"]
    imports = bal["imports"]
    balance = bal["balance"]

    plt.rcParams["font.family"] = chart_font()
    fig, ax = plt.subplots(figsize=(7.6, 4.4), dpi=160)
    x = list(range(len(years)))
    w = 0.26
    cols = [lighten(THEME_ACCENTS[0], 0.30), lighten(THEME_ACCENTS[1], 0.30),
            lighten(THEME_ACCENTS[2], 0.30)]
    ax.bar([i - w for i in x], exports, width=w, label="Exports", color=cols[0])
    ax.bar([i for i in x], imports, width=w, label="Imports", color=cols[1])
    ax.bar([i + w for i in x], balance, width=w, label="Balance of Trade", color=cols[2])
    ax.set_xticks(x)
    ax.set_xticklabels([str(y) for y in years], fontsize=9)
    ax.set_ylabel("Value in USD Million", fontsize=10)
    ax.set_title(f"Kenya Services Balance of Trade (USD Million)", fontsize=12, weight="bold")
    ax.yaxis.grid(True, linestyle="--", alpha=0.35)
    ax.set_axisbelow(True)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.10), ncol=3, frameon=False, fontsize=10)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)


def make_chart_export_share(a: ServicesAnalysis, out_path):
    """Doughnut chart of Kenya's service export shares by category."""
    items = a.table3["items"][:10]
    labels = [short_label(d["label"] or d["name"]) for d in items]
    shares = [d["share"] for d in items if d["share"] is not None]
    labels = labels[:len(shares)]
    other = max(0.0, 1.0 - sum(shares))
    if other > 0.01:
        labels.append("Other services")
        shares.append(other)

    palette = list(THEME_ACCENTS) + [lighten(c, 0.45) for c in THEME_ACCENTS]
    colors = palette[:len(shares)]

    plt.rcParams["font.family"] = chart_font()
    fig, ax = plt.subplots(figsize=(7.8, 4.6), dpi=160)
    wedges, _, autotexts = ax.pie(
        shares, labels=None, autopct="%1.1f%%", startangle=90,
        counterclock=False, colors=colors, pctdistance=0.80,
        wedgeprops=dict(width=0.42, edgecolor="white", linewidth=1.2))
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color("white")
    ax.set_title(f"Share of Kenya's Service Exports by Category in {a.year}",
                 fontsize=12, weight="bold")
    ax.legend(wedges, [f"{l} - {s * 100:.1f}%" for l, s in zip(labels, shares)],
              loc="upper center", bbox_to_anchor=(0.5, -0.03), ncol=2,
              frameon=False, fontsize=8.5)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)


def make_chart_import_share(a: ServicesAnalysis, out_path):
    """Doughnut chart of Kenya's service import shares by category."""
    items = a.table4["items"][:10]
    labels = [short_label(d["label"] or d["name"]) for d in items]
    shares = [d["share"] for d in items if d["share"] is not None]
    labels = labels[:len(shares)]
    other = max(0.0, 1.0 - sum(shares))
    if other > 0.01:
        labels.append("Other services")
        shares.append(other)

    palette = list(THEME_ACCENTS) + [lighten(c, 0.45) for c in THEME_ACCENTS]
    colors = palette[:len(shares)]

    plt.rcParams["font.family"] = chart_font()
    fig, ax = plt.subplots(figsize=(7.8, 4.6), dpi=160)
    wedges, _, autotexts = ax.pie(
        shares, labels=None, autopct="%1.1f%%", startangle=90,
        counterclock=False, colors=colors, pctdistance=0.80,
        wedgeprops=dict(width=0.42, edgecolor="white", linewidth=1.2))
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color("white")
    ax.set_title(f"Share of Kenya's Service Imports by Category in {a.year}",
                 fontsize=12, weight="bold")
    ax.legend(wedges, [f"{l} - {s * 100:.1f}%" for l, s in zip(labels, shares)],
              loc="upper center", bbox_to_anchor=(0.5, -0.03), ncol=2,
              frameon=False, fontsize=8.5)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)


# ---------------------------------------------------------------------------
# Word document building
# ---------------------------------------------------------------------------
def build_services_report(cfg, excel_dir, out_path, tmp_dir):
    os.makedirs(tmp_dir, exist_ok=True)
    a = ServicesAnalysis(cfg).load(excel_dir)
    narr = build_narratives(a, cfg)

    c = cfg["country"]
    rep = cfg["report"]
    Y = a.year

    b = ReportBuilder(cfg, narr)
    b.a = a
    doc = b.doc

    # ============================== TITLE PAGE ==============================
    svc_title = f"KENYA- {c['title']} SERVICES TRADE FLOW"
    b.add_para(svc_title,
               size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    b.add_para(rep.get("title_line2", "ANALYSIS REPORT"),
               size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    for _ in range(3):
        b.add_para("")
    b.add_para("Prepared", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    b.add_para("BY", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    b.add_para("Kenya Export Promotion and Branding Agency (KEPROBA)",
               size=22, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    b.add_para("Research and Innovation Directorate (RID)",
               size=22, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    b.add_para(rep.get("month_year", ""), size=22, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    b.page_break()

    # ============================== TABLE OF CONTENTS =======================
    b.add_toc()
    b.page_break()

    # ============================== SECTION 1 ===============================
    b.add_heading(f"1. {c['title']}")
    b.add_heading("1.1. Backgrounds", level=2)
    for segments in narr["background"]:
        b.add_background_para(segments)
    b.page_break()

    # ============================== SECTION 2 ===============================
    b.add_heading("2. Global Services Trade")

    # Table 1: top exporters
    n1 = len(a.table1["items"])
    b.add_table_caption(f"Table 1: Top {min(n1, 20)} Global Service Exporters")
    b.add_market_table(a, a.table1, importer=False, unit_row=True)
    b.add_source()
    for line in narr.get("s2_exporters", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # Export share chart
    b.add_para("Share of Global Service Exports by Country", bold=True)
    b.add_table_caption(f"Figure 1: Share of Global Service Exports in {Y}")
    fig1 = os.path.join(tmp_dir, "chart_export_share.png")
    make_chart_export_share(a, fig1)
    b.add_figure(fig1)
    b.add_source()
    b.page_break()

    # ============================== SECTION 3 ===============================
    b.add_heading(f"3. Top Service Importers in {Y}")

    # Table 2: top importers
    n2 = len(a.table2["items"])
    b.add_table_caption(f"Table 2: Top {min(n2, 20)} Global Service Importers")
    b.add_market_table(a, a.table2, importer=True, unit_row=True)
    b.add_source()
    for line in narr.get("s3_importers", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # ============================== SECTION 3.1 =============================
    b.add_heading(f"3.1 Kenya's Service Exports")

    # Table 3: Kenya service exports
    b.add_table_caption(f"Table 3: Kenya's Service Exports by Category in {Y}")
    b.add_product_table(a, a.table3,
                        f"Kenya's Service Exports\nValue in USD Million",
                        unit_row=True,
                        widths=[685, 972, 3170, 731, 731, 731, 731, 733, 866])
    b.add_source()
    for line in narr.get("s3_kenya_exports", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # Import share chart
    b.add_para("Share of Kenya's Service Imports by Category", bold=True)
    b.add_table_caption(f"Figure 2: Share of Kenya's Service Imports in {Y}")
    fig2 = os.path.join(tmp_dir, "chart_import_share.png")
    make_chart_import_share(a, fig2)
    b.add_figure(fig2)
    b.add_source()
    b.page_break()

    # ============================== SECTION 3.2 =============================
    b.add_heading("3.2 Service Exporters by Development Status")

    # Table 5: Exporters by development status
    b.add_table_caption(f"Table 5: Top Global Service Exporters by Development Status in {Y}")
    t5_path = os.path.join(excel_dir, "Table 5 Service Exporters by Dev Status.xlsx")
    if os.path.exists(t5_path):
        b.add_table_from_excel(t5_path, "Table 5")
    else:
        b.add_para("[Table 5 data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.add_source()
    b.add_para(
        "Service exporters are classified as Developed (high-income economies), "
        "Developing (middle-income economies), or Least Developed Countries (LDCs) "
        "based on World Bank and UNCTAD classifications.",
        italic=True)
    b.page_break()

    # ============================== SECTION 3.3 =============================
    b.add_heading("3.3 Service Importers by Development Status")

    # Table 6: Importers by development status
    b.add_table_caption(f"Table 6: Top Global Service Importers by Development Status in {Y}")
    t6_path = os.path.join(excel_dir, "Table 6 Service Importers by Dev Status.xlsx")
    if os.path.exists(t6_path):
        b.add_table_from_excel(t6_path, "Table 6")
    else:
        b.add_para("[Table 6 data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.add_source()
    b.page_break()

    # ============================== SECTION 3.4 =============================
    b.add_heading("3.4 Structure of Global Service Exports")

    # Pie chart of service exports structure
    pie_path = os.path.join(excel_dir, "Figure 2 Service Exports Structure.png")
    if os.path.exists(pie_path):
        b.add_table_caption("Figure 4: Structure of Global Service Exports by Category")
        b.add_figure(pie_path)
        b.add_source()
    else:
        b.add_para("[Pie chart data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))

    # Stacked bar chart
    bar_path = os.path.join(excel_dir, "Figure 3 Service Exports by Category.png")
    if os.path.exists(bar_path):
        b.add_table_caption("Figure 5: Global Service Exports by Category Over Time")
        b.add_figure(bar_path)
        b.add_source()
    else:
        b.add_para("[Stacked bar chart data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.page_break()

    # ============================== SECTION 3.5 =============================
    b.add_heading("3.5 Regional Analysis of Service Trade")

    # Table 7: Exporters by region
    b.add_table_caption(f"Table 7: Top Service Exporters by Region in {Y}")
    t7_path = os.path.join(excel_dir, "Table 7 Service Exporters by Region.xlsx")
    if os.path.exists(t7_path):
        b.add_table_from_excel(t7_path, "Table 7")
    else:
        b.add_para("[Table 7 data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.add_source()
    b.add_para(
        "Service exporters are grouped by ITC regional classifications: "
        "Africa, Asia and the Pacific, Eastern Europe and Central Asia, "
        "Latin America and the Caribbean, and Middle East and North Africa.",
        italic=True)
    b.page_break()

    # Regional exporter chart
    reg_exp_path = os.path.join(excel_dir, "Figure 4 Service Exports by Region.png")
    if os.path.exists(reg_exp_path):
        b.add_table_caption("Figure 6: Top Service Exporters by Region")
        b.add_figure(reg_exp_path)
        b.add_source()
    b.page_break()

    # Table 8: Importers by region
    b.add_table_caption(f"Table 8: Top Service Importers by Region in {Y}")
    t8_path = os.path.join(excel_dir, "Table 8 Service Importers by Region.xlsx")
    if os.path.exists(t8_path):
        b.add_table_from_excel(t8_path, "Table 8")
    else:
        b.add_para("[Table 8 data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.add_source()
    b.page_break()

    # Regional importer chart
    reg_imp_path = os.path.join(excel_dir, "Figure 5 Service Imports by Region.png")
    if os.path.exists(reg_imp_path):
        b.add_table_caption("Figure 7: Top Service Importers by Region")
        b.add_figure(reg_imp_path)
        b.add_source()
    b.page_break()

    # ============================== SECTION 3.6 =============================
    b.add_heading("3.6 Kenya's Revealed Comparative Advantage in Services")

    # Table 9: Kenya's RCA
    b.add_table_caption(f"Table 9: Kenya's Revealed Comparative Advantage (RCA) in Services ({Y})")
    t9_path = os.path.join(excel_dir, "Table 9 Kenya Services RCA.xlsx")
    if os.path.exists(t9_path):
        b.add_table_from_excel(t9_path, "Table 9")
    else:
        b.add_para("[Table 9 data not available]", italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.add_source()
    b.add_para(
        "Revealed Comparative Advantage (RCA) is computed using the Balassa Index: "
        "RCA = (Kenya's share of exports in category i) / (World's share of exports in category i). "
        "RCA > 2.5 indicates strong comparative advantage; RCA 1.0-2.5 indicates moderate advantage; "
        "RCA < 1.0 indicates comparative disadvantage.",
        italic=True)
    b.page_break()

    # Figure 6: RCA chart
    rca_chart = os.path.join(excel_dir, "Figure 6 Kenya Services RCA.png")
    if os.path.exists(rca_chart):
        b.add_table_caption("Figure 8: Kenya's Service Export RCA by Category")
        b.add_figure(rca_chart)
        b.add_source()
    b.page_break()

    # ============================== SECTION 4 ===============================
    b.add_heading(f"4. {c['name']}'s Services Sector")
    b.add_para(
        f"{c['name']} has a growing services sector that contributes "
        f"significantly to the country's GDP.", italic=True)
    b.add_source("Source: [RESEARCH NEEDED: GDP contribution data for services sector]")

    b.add_heading("4.1 Services and Employment", level=2)
    b.add_para("[RESEARCH NEEDED: Employment data in the services sector, "
               "including key sub-sectors and employment intensity]",
               italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))

    b.add_heading("4.2 Integrated National Export Development and Promotion Strategy", level=2)
    b.add_para("[RESEARCH NEEDED: INEDPS strategy details for services, "
               "including priority service categories and export targets]",
               italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))

    b.add_heading("4.3 Key Service Sector Projections", level=2)
    for sector in ["Financial Services", "Tourism", "ICT", "Transport",
                    "Business Process Outsourcing", "Education"]:
        b.add_heading(f"4.3.x {sector}", level=2)
        b.add_para(f"[RESEARCH NEEDED: {sector} sector overview, projections, "
                   f"and market data for {c['name']}]",
                   italic=True, color=RGBColor(0x9A, 0x1F, 0x1F))
    b.page_break()

    # ============================== SECTION 5 ===============================
    b.add_heading("5. Kenya's Services Trade Balance")

    # Figure 1: Balance chart
    b.add_table_caption(f"Figure 3: Kenya's Services Balance of Trade")
    fig3 = os.path.join(tmp_dir, "chart_balance.png")
    make_chart_balance(a, fig3)
    b.add_figure(fig3)
    b.add_source()
    for line in narr.get("s5", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # Table 4: Kenya service imports
    b.add_heading(f"5.1 Kenya's Service Imports", level=2)
    b.add_table_caption(f"Table 4: Kenya's Service Imports by Category in {Y}")
    b.add_product_table(a, a.table4,
                        f"Kenya's Service Imports\nValue in USD Million",
                        unit_row=True,
                        widths=[685, 972, 3170, 731, 731, 731, 731, 733, 866])
    b.add_source()
    for line in narr.get("s4_imports", []):
        if line:
            b.add_bullet(line)
    b.page_break()

    # ============================== ANNEX ===================================
    b.add_heading("6. Annex: Quick Facts")
    b.add_quick_facts(cfg.get("quick_facts", []))
    b.page_break()

    # ============================== REFERENCES ==============================
    b.add_heading("References")
    for ref in cfg.get("references", []):
        b.add_para(ref)

    b.add_footer()
    doc.save(out_path)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(
        description="Generate a Kenya - Country Services Trade Flow report.")
    ap.add_argument("--excel-dir", default=os.path.join(BASE_DIR, "output", "service_tables"),
                    help="Folder containing the five service Excel files")
    ap.add_argument("--config", default=os.path.join(BASE_DIR, "config", "saudi_arabia.json"),
                    help="Per-country JSON configuration file")
    ap.add_argument("--output", default=None, help="Output .docx path")
    ap.add_argument("--tmp", default=os.path.join(BASE_DIR, "output", ".tmp"),
                    help="Temporary directory for chart images")
    args = ap.parse_args()

    global cfg_path
    cfg_path = os.path.abspath(args.config)
    cfg = load_config(cfg_path)

    out = args.output
    if not out:
        cname = cfg["country"]["name"].replace(" ", "_")
        out = os.path.join(BASE_DIR, "output",
                           f"KENYA-{cname} SERVICES TRADE FLOW.docx")
    out = os.path.abspath(out)
    out_dir = os.path.dirname(out)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    os.makedirs(args.tmp, exist_ok=True)

    print(f"[1/3] Reading service Excel data from : {os.path.abspath(args.excel_dir)}")
    print(f"[2/3] Building services report for    : {cfg['country']['name']}")
    build_services_report(cfg, args.excel_dir, out, args.tmp)
    print(f"[3/3] Report saved to                 : {out}")


if __name__ == "__main__":
    main()
